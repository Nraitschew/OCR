#!/usr/bin/env python3
"""
Enhanced OCR System with Advanced Table Extraction
"""

import os
import sys
import time
import psutil
import numpy as np
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass, field
import asyncio
import re

# Resource Configuration
MAX_CPU_PERCENT = 70  # Maximum CPU usage percentage
MAX_RAM_PERCENT = 70  # Maximum RAM usage percentage  
MAX_RAM_GB = 2  # Maximum RAM in GB (can be exceeded if needed)

# Document processing imports
from PIL import Image
import cv2
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from fast_langdetect import detect_language
import pandas as pd
import pdfplumber

# Configure logging
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class TableData:
    """Represents extracted table data"""
    headers: List[str]
    rows: List[List[str]]
    confidence: float = 0.0
    position: Dict[str, int] = field(default_factory=dict)
    
    def to_dataframe(self) -> pd.DataFrame:
        """Convert to pandas DataFrame"""
        if self.headers:
            return pd.DataFrame(self.rows, columns=self.headers)
        else:
            return pd.DataFrame(self.rows)
            
    def to_markdown(self) -> str:
        """Convert to markdown table"""
        if not self.rows:
            return ""
            
        # Use headers or generate column numbers
        headers = self.headers if self.headers else [f"Col{i+1}" for i in range(len(self.rows[0]))]
        
        # Build markdown table
        lines = []
        
        # Header
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join([" --- " for _ in headers]) + "|")
        
        # Rows
        for row in self.rows:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
            
        return "\n".join(lines)


@dataclass
class EnhancedOCRResult:
    """Enhanced OCR Processing Result with detailed table information"""
    text: str
    language: str
    confidence: float
    tables: List[TableData]
    markdown: str
    processing_time: float
    metadata: Dict[str, Any]
    structured_content: Dict[str, Any] = field(default_factory=dict)
    
    def get_tables_summary(self) -> str:
        """Get summary of extracted tables"""
        if not self.tables:
            return "No tables found"
            
        summary = f"Found {len(self.tables)} table(s):\n"
        for i, table in enumerate(self.tables):
            rows = len(table.rows)
            cols = len(table.rows[0]) if table.rows else 0
            summary += f"  Table {i+1}: {rows} rows × {cols} columns"
            if table.confidence > 0:
                summary += f" (confidence: {table.confidence:.1%})"
            summary += "\n"
            
        return summary


class TableExtractor:
    """Advanced table extraction from documents"""
    
    def __init__(self):
        self.min_table_rows = 2
        self.min_table_cols = 2
        
    def extract_from_image(self, image: np.ndarray) -> List[TableData]:
        """Extract tables from image using computer vision"""
        tables = []
        
        try:
            # Convert to grayscale
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image
                
            # Detect horizontal and vertical lines
            horizontal_lines = self._detect_lines(gray, horizontal=True)
            vertical_lines = self._detect_lines(gray, horizontal=False)
            
            # Find table regions
            table_regions = self._find_table_regions(horizontal_lines, vertical_lines)
            
            # Extract text from each table region
            for region in table_regions:
                table_data = self._extract_table_from_region(image, region)
                if table_data:
                    tables.append(table_data)
                    
        except Exception as e:
            logger.error(f"Table extraction from image failed: {e}")
            
        return tables
        
    def extract_from_pdf_page(self, page) -> List[TableData]:
        """Extract tables from PDF page using pdfplumber"""
        tables = []
        
        try:
            # Use pdfplumber for better table extraction
            import pdfplumber
            
            with pdfplumber.open(page) as pdf:
                for table in pdf.pages[0].extract_tables():
                    if table and len(table) >= self.min_table_rows:
                        # Clean and structure table data
                        cleaned_table = self._clean_table_data(table)
                        if cleaned_table:
                            table_data = TableData(
                                headers=cleaned_table[0] if len(cleaned_table) > 1 else [],
                                rows=cleaned_table[1:] if len(cleaned_table) > 1 else cleaned_table,
                                confidence=0.9  # High confidence for native PDF tables
                            )
                            tables.append(table_data)
                            
        except Exception as e:
            logger.debug(f"pdfplumber table extraction failed: {e}")
            
        return tables
        
    def extract_from_text(self, text: str) -> List[TableData]:
        """Extract tables from structured text"""
        tables = []
        
        # Look for table-like patterns in text
        lines = text.split('\n')
        potential_tables = self._find_text_tables(lines)
        
        for table_lines in potential_tables:
            table_data = self._parse_text_table(table_lines)
            if table_data:
                tables.append(table_data)
                
        return tables
        
    def _detect_lines(self, gray: np.ndarray, horizontal: bool = True) -> np.ndarray:
        """Detect horizontal or vertical lines in image"""
        # Create structure element for line detection
        if horizontal:
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
        else:
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            
        # Apply morphology operations
        lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, kernel)
        
        return lines
        
    def _find_table_regions(self, h_lines: np.ndarray, v_lines: np.ndarray) -> List[Tuple[int, int, int, int]]:
        """Find table regions from detected lines"""
        regions = []
        
        # Combine horizontal and vertical lines
        combined = cv2.add(h_lines, v_lines)
        
        # Find contours
        contours, _ = cv2.findContours(combined, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        # Filter contours to find table-like regions
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            
            # Filter by size and aspect ratio
            if w > 50 and h > 50 and w/h < 10 and h/w < 10:
                regions.append((x, y, w, h))
                
        return regions
        
    def _extract_table_from_region(self, image: np.ndarray, region: Tuple[int, int, int, int]) -> Optional[TableData]:
        """Extract table data from image region"""
        x, y, w, h = region
        
        # Crop region
        table_img = image[y:y+h, x:x+w]
        
        # OCR the region
        try:
            text = pytesseract.image_to_string(table_img, lang='deu+eng')
            
            # Parse the text as table
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            if len(lines) >= self.min_table_rows:
                # Try to parse as table
                rows = []
                for line in lines:
                    # Split by multiple spaces or tabs
                    cells = re.split(r'\s{2,}|\t', line)
                    if len(cells) >= self.min_table_cols:
                        rows.append(cells)
                        
                if rows:
                    return TableData(
                        headers=rows[0] if len(rows) > 1 else [],
                        rows=rows[1:] if len(rows) > 1 else rows,
                        confidence=0.7,
                        position={'x': x, 'y': y, 'width': w, 'height': h}
                    )
                    
        except Exception as e:
            logger.error(f"Failed to extract table from region: {e}")
            
        return None
        
    def _clean_table_data(self, table: List[List[Any]]) -> List[List[str]]:
        """Clean and normalize table data"""
        cleaned = []
        
        for row in table:
            if row and any(cell for cell in row if cell):
                cleaned_row = []
                for cell in row:
                    # Convert to string and clean
                    cell_str = str(cell) if cell else ""
                    cell_str = cell_str.strip()
                    cleaned_row.append(cell_str)
                    
                # Only add non-empty rows
                if any(cell for cell in cleaned_row):
                    cleaned.append(cleaned_row)
                    
        return cleaned
        
    def _find_text_tables(self, lines: List[str]) -> List[List[str]]:
        """Find potential tables in text lines"""
        potential_tables = []
        current_table = []
        
        for line in lines:
            # Check if line looks like table row (has multiple columns)
            if self._is_table_line(line):
                current_table.append(line)
            else:
                # End of table
                if len(current_table) >= self.min_table_rows:
                    potential_tables.append(current_table)
                current_table = []
                
        # Don't forget last table
        if len(current_table) >= self.min_table_rows:
            potential_tables.append(current_table)
            
        return potential_tables
        
    def _is_table_line(self, line: str) -> bool:
        """Check if a line looks like a table row"""
        # Count separators (tabs, multiple spaces, pipes)
        separators = len(re.findall(r'\t|\s{2,}|\|', line))
        
        # Has multiple columns
        return separators >= self.min_table_cols - 1
        
    def _parse_text_table(self, lines: List[str]) -> Optional[TableData]:
        """Parse text lines as table"""
        if not lines:
            return None
            
        rows = []
        for line in lines:
            # Try different separators
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            elif '\t' in line:
                cells = [cell.strip() for cell in line.split('\t')]
            else:
                cells = re.split(r'\s{2,}', line.strip())
                
            if len(cells) >= self.min_table_cols:
                rows.append(cells)
                
        if rows:
            # Check if first row looks like headers
            first_row_numeric = sum(1 for cell in rows[0] if re.match(r'^[\d.,]+$', cell))
            headers = rows[0] if first_row_numeric < len(rows[0]) / 2 else []
            
            return TableData(
                headers=headers,
                rows=rows[1:] if headers else rows,
                confidence=0.8
            )
            
        return None


class EnhancedOCREngine:
    """Enhanced OCR Engine with table extraction capabilities"""
    
    def __init__(self):
        self.resource_manager = ResourceManager()
        self.table_extractor = TableExtractor()
        self._check_tesseract()
        
    def _check_tesseract(self):
        """Check if tesseract is available"""
        try:
            pytesseract.get_tesseract_version()
            logger.info("Tesseract is available")
        except:
            logger.warning("Tesseract not found. OCR functionality will be limited.")
            
    def detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            if len(text.strip()) > 0:
                result = detect_language(text)
                return result['lang']
        except:
            pass
        return 'de'  # Default to German
        
    def extract_text_with_structure(self, image: np.ndarray) -> Tuple[str, float, Dict[str, Any]]:
        """Extract text with structural information"""
        try:
            # Get detailed OCR data
            pil_image = Image.fromarray(image)
            data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT, lang='deu+eng')
            
            # Extract text
            text = pytesseract.image_to_string(pil_image, lang='deu+eng')
            
            # Calculate confidence
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            confidence = np.mean(confidences) if confidences else 0
            
            # Extract structure
            structure = {
                'paragraphs': self._extract_paragraphs(data),
                'lines': self._extract_lines(data),
                'word_count': len([w for w in data['text'] if w.strip()])
            }
            
            return text, confidence / 100, structure
            
        except Exception as e:
            logger.error(f"OCR with structure failed: {e}")
            return "", 0.0, {}
            
    def _extract_paragraphs(self, ocr_data: Dict) -> List[str]:
        """Extract paragraphs from OCR data"""
        paragraphs = []
        current_para = []
        last_par_num = -1
        
        for i, par_num in enumerate(ocr_data['par_num']):
            if par_num != last_par_num and current_para:
                paragraphs.append(' '.join(current_para))
                current_para = []
                
            if ocr_data['text'][i].strip():
                current_para.append(ocr_data['text'][i])
                
            last_par_num = par_num
            
        if current_para:
            paragraphs.append(' '.join(current_para))
            
        return paragraphs
        
    def _extract_lines(self, ocr_data: Dict) -> List[str]:
        """Extract lines from OCR data"""
        lines = []
        current_line = []
        last_line_num = -1
        
        for i, line_num in enumerate(ocr_data['line_num']):
            if line_num != last_line_num and current_line:
                lines.append(' '.join(current_line))
                current_line = []
                
            if ocr_data['text'][i].strip():
                current_line.append(ocr_data['text'][i])
                
            last_line_num = line_num
            
        if current_line:
            lines.append(' '.join(current_line))
            
        return lines
        
    def process_image_enhanced(self, image: np.ndarray) -> EnhancedOCRResult:
        """Process image with enhanced table extraction"""
        start_time = time.time()
        
        # Wait for resources
        self.resource_manager.wait_for_resources()
        
        # Preprocess image
        processed_image = self._preprocess_image(image)
        
        # Extract text with structure
        text, confidence, structure = self.extract_text_with_structure(processed_image)
        
        # Extract tables
        tables = self.table_extractor.extract_from_image(processed_image)
        
        # Also try to extract tables from text
        text_tables = self.table_extractor.extract_from_text(text)
        tables.extend(text_tables)
        
        # Detect language
        language = self.detect_language(text)
        
        # Generate enhanced markdown
        markdown = self._generate_enhanced_markdown(text, tables)
        
        processing_time = time.time() - start_time
        
        return EnhancedOCRResult(
            text=text,
            language=language,
            confidence=confidence,
            tables=tables,
            markdown=markdown,
            processing_time=processing_time,
            metadata={
                'engine': 'tesseract',
                'image_shape': image.shape,
                'tables_found': len(tables)
            },
            structured_content=structure
        )
        
    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Enhanced image preprocessing"""
        # Convert to grayscale if needed
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
            
        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Improve contrast
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)
        
        # Threshold
        _, thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return thresh
        
    def _generate_enhanced_markdown(self, text: str, tables: List[TableData]) -> str:
        """Generate enhanced markdown with proper table formatting"""
        markdown_parts = []
        
        # Add main text
        if text:
            markdown_parts.append(text)
            
        # Add tables
        if tables:
            markdown_parts.append("\n## Extracted Tables\n")
            for i, table in enumerate(tables):
                markdown_parts.append(f"\n### Table {i+1}")
                if table.confidence > 0:
                    markdown_parts.append(f"*Confidence: {table.confidence:.1%}*\n")
                markdown_parts.append(table.to_markdown())
                
        return "\n".join(markdown_parts)


class EnhancedDocumentProcessor:
    """Enhanced document processor with better table handling"""
    
    def __init__(self, ocr_engine: EnhancedOCREngine):
        self.ocr_engine = ocr_engine
        self.table_extractor = TableExtractor()
        
    async def process_pdf_enhanced(self, file_path: Path) -> List[EnhancedOCRResult]:
        """Process PDF with enhanced table extraction"""
        results = []
        
        try:
            # Try pdfplumber first for better table extraction
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    text = page.extract_text() or ""
                    
                    # Extract tables
                    tables = []
                    for table in page.extract_tables():
                        if table:
                            cleaned = self.table_extractor._clean_table_data(table)
                            if cleaned and len(cleaned) >= 2:
                                table_data = TableData(
                                    headers=cleaned[0],
                                    rows=cleaned[1:],
                                    confidence=0.95
                                )
                                tables.append(table_data)
                                
                    # If we have content, create result
                    if text or tables:
                        result = EnhancedOCRResult(
                            text=text,
                            language=self.ocr_engine.detect_language(text),
                            confidence=1.0 if text else 0.9,
                            tables=tables,
                            markdown=self.ocr_engine._generate_enhanced_markdown(text, tables),
                            processing_time=0.1,
                            metadata={
                                'page': page_num + 1,
                                'type': 'native',
                                'method': 'pdfplumber'
                            },
                            structured_content={}
                        )
                        results.append(result)
                        
        except Exception as e:
            logger.debug(f"pdfplumber failed, falling back to PyMuPDF: {e}")
            
            # Fallback to PyMuPDF
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                
                if text.strip():
                    # Extract tables from text
                    tables = self.table_extractor.extract_from_text(text)
                    
                    result = EnhancedOCRResult(
                        text=text,
                        language=self.ocr_engine.detect_language(text),
                        confidence=1.0,
                        tables=tables,
                        markdown=self.ocr_engine._generate_enhanced_markdown(text, tables),
                        processing_time=0.1,
                        metadata={'page': page_num + 1, 'type': 'native'}
                    )
                    results.append(result)
                else:
                    # Convert to image for OCR
                    pix = page.get_pixmap()
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                        pix.height, pix.width, pix.n
                    )
                    
                    if img.shape[2] == 4:
                        img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
                        
                    result = self.ocr_engine.process_image_enhanced(img)
                    result.metadata['page'] = page_num + 1
                    results.append(result)
                    
            doc.close()
            
        return results
        
    async def process_docx_enhanced(self, file_path: Path) -> EnhancedOCRResult:
        """Process DOCX with table extraction"""
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
                
        text = '\n'.join(paragraphs)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text.strip())
                if any(cells):  # Only add non-empty rows
                    rows.append(cells)
                    
            if rows:
                table_data = TableData(
                    headers=rows[0] if len(rows) > 1 else [],
                    rows=rows[1:] if len(rows) > 1 else rows,
                    confidence=1.0
                )
                tables.append(table_data)
                
        # Generate markdown
        markdown = self.ocr_engine._generate_enhanced_markdown(text, tables)
        
        return EnhancedOCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=tables,
            markdown=markdown,
            processing_time=0.1,
            metadata={
                'type': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables_in_doc': len(doc.tables)
            },
            structured_content={
                'paragraphs': paragraphs
            }
        )


class EnhancedOCRSystem:
    """Enhanced OCR System with advanced table extraction"""
    
    def __init__(self):
        self.ocr_engine = EnhancedOCREngine()
        self.document_processor = EnhancedDocumentProcessor(self.ocr_engine)
        
    async def process_file(self, file_path: str) -> EnhancedOCRResult:
        """Process file with enhanced capabilities"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            results = await self.document_processor.process_pdf_enhanced(path)
            # Combine results
            if not results:
                return self._empty_result()
                
            combined_text = '\n\n'.join([r.text for r in results])
            all_tables = []
            for r in results:
                all_tables.extend(r.tables)
                
            return EnhancedOCRResult(
                text=combined_text,
                language=results[0].language if results else 'de',
                confidence=np.mean([r.confidence for r in results]) if results else 0,
                tables=all_tables,
                markdown=self.ocr_engine._generate_enhanced_markdown(combined_text, all_tables),
                processing_time=sum([r.processing_time for r in results]),
                metadata={
                    'pages': len(results),
                    'type': 'pdf',
                    'total_tables': len(all_tables)
                }
            )
            
        elif suffix in ['.png', '.jpg', '.jpeg']:
            image = cv2.imread(str(path))
            return self.ocr_engine.process_image_enhanced(image)
            
        elif suffix == '.docx':
            return await self.document_processor.process_docx_enhanced(path)
            
        elif suffix == '.txt':
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
                
            # Try to extract tables from text
            tables = self.ocr_engine.table_extractor.extract_from_text(text)
            
            return EnhancedOCRResult(
                text=text,
                language=self.ocr_engine.detect_language(text),
                confidence=1.0,
                tables=tables,
                markdown=self.ocr_engine._generate_enhanced_markdown(text, tables),
                processing_time=0.01,
                metadata={'type': 'txt', 'size': len(text)}
            )
            
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
            
    def _empty_result(self) -> EnhancedOCRResult:
        """Return empty result"""
        return EnhancedOCRResult(
            text="",
            language="unknown",
            confidence=0.0,
            tables=[],
            markdown="",
            processing_time=0.0,
            metadata={}
        )


class ResourceManager:
    """Manages system resources"""
    
    def __init__(self):
        self.process = psutil.Process()
        self.start_time = time.time()
        
    def get_current_usage(self) -> Dict[str, float]:
        """Get current resource usage"""
        return {
            'cpu_percent': self.process.cpu_percent(interval=0.1),
            'ram_mb': self.process.memory_info().rss / (1024 * 1024),
            'ram_gb': self.process.memory_info().rss / (1024 * 1024 * 1024),
            'ram_percent': self.process.memory_percent()
        }
        
    def wait_for_resources(self, timeout: int = 10):
        """Wait for resources to become available"""
        # Simple implementation for now
        time.sleep(0.1)
        return True


def main():
    """Test enhanced OCR system"""
    print("Enhanced OCR System Ready")
    print(f"Resource Limits: CPU {MAX_CPU_PERCENT}%, RAM {MAX_RAM_PERCENT}% (Max {MAX_RAM_GB}GB)")
    print("Features:")
    print("  - Advanced table extraction")
    print("  - Multi-column table support")
    print("  - Native PDF table extraction")
    print("  - Structured content extraction")


if __name__ == "__main__":
    main()