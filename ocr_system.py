#!/usr/bin/env python3
"""
Comprehensive OCR System with Resource Management
Supports: PDF, DOCX, PNG, JPEG, TXT
"""

import os
import sys
import time
import asyncio
import logging
import psutil
import numpy as np
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
from dataclasses import dataclass
import multiprocessing as mp
import resource

# Resource Configuration - EASILY CONFIGURABLE
MAX_CPU_PERCENT = 70  # Maximum CPU usage percentage
MAX_RAM_PERCENT = 70  # Maximum RAM usage percentage  
MAX_RAM_GB = 2  # Maximum RAM in GB (can be exceeded if needed)

# Document processing imports
from PIL import Image
import cv2
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from docx import Document
import pdfplumber
from tabulate import tabulate

# OCR imports
import easyocr
import pytesseract
from paddleocr import PaddleOCR

# Language detection
from fast_langdetect import detect_language

# Utilities
import pandas as pd
import torch
import requests
from bs4 import BeautifulSoup

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class OCRConfig:
    """OCR System Configuration"""
    use_gpu: bool = torch.cuda.is_available()
    languages: List[str] = None
    max_workers: int = mp.cpu_count()
    batch_size: int = 10
    timeout: int = 300
    
    def __post_init__(self):
        if self.languages is None:
            self.languages = ['en', 'de']


@dataclass
class OCRResult:
    """OCR Processing Result"""
    text: str
    language: str
    confidence: float
    tables: List[pd.DataFrame]
    markdown: str
    processing_time: float
    metadata: Dict[str, Any]


class ResourceManager:
    """Manages system resources to stay within limits"""
    
    def __init__(self):
        self.process = psutil.Process()
        self.start_time = time.time()
        
    def get_current_usage(self) -> Dict[str, float]:
        """Get current resource usage"""
        return {
            'cpu_percent': self.process.cpu_percent(interval=0.1),
            'ram_mb': self.process.memory_info().rss / (1024 * 1024),
            'ram_percent': self.process.memory_percent()
        }
    
    def check_limits(self) -> bool:
        """Check if we're within resource limits"""
        usage = self.get_current_usage()
        
        if usage['cpu_percent'] > MAX_CPU_PERCENT:
            logger.warning(f"CPU usage {usage['cpu_percent']:.1f}% exceeds limit {MAX_CPU_PERCENT}%")
            return False
            
        if usage['ram_percent'] > MAX_RAM_PERCENT:
            logger.warning(f"RAM usage {usage['ram_percent']:.1f}% exceeds limit {MAX_RAM_PERCENT}%")
            # Allow exceeding if necessary
            
        return True
    
    def wait_for_resources(self, timeout: int = 10):
        """Wait for resources to become available"""
        start = time.time()
        while time.time() - start < timeout:
            if self.check_limits():
                return True
            time.sleep(0.5)
        return False


class OCREngine:
    """Main OCR Engine with GPU/CPU fallback"""
    
    def __init__(self, config: OCRConfig):
        self.config = config
        self.resource_manager = ResourceManager()
        self._initialize_engines()
        
    def _initialize_engines(self):
        """Initialize OCR engines based on availability"""
        self.engines = {}
        
        # Initialize EasyOCR
        try:
            self.engines['easyocr'] = easyocr.Reader(
                self.config.languages, 
                gpu=self.config.use_gpu
            )
            logger.info(f"EasyOCR initialized (GPU: {self.config.use_gpu})")
        except Exception as e:
            logger.error(f"Failed to initialize EasyOCR: {e}")
            
        # Initialize PaddleOCR for tables
        try:
            self.engines['paddle'] = PaddleOCR(
                use_angle_cls=True,
                lang='en',
                use_gpu=self.config.use_gpu
            )
            logger.info("PaddleOCR initialized")
        except Exception as e:
            logger.error(f"Failed to initialize PaddleOCR: {e}")
            
        # Tesseract is always available as fallback
        self.engines['tesseract'] = pytesseract
        logger.info("Tesseract initialized as fallback")
        
    def detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            result = detect_language(text)
            return result['lang']
        except:
            return 'en'
            
    def extract_text_easyocr(self, image: np.ndarray) -> Tuple[str, float]:
        """Extract text using EasyOCR"""
        if 'easyocr' not in self.engines:
            raise ValueError("EasyOCR not available")
            
        results = self.engines['easyocr'].readtext(image)
        text = ' '.join([item[1] for item in results])
        confidence = np.mean([item[2] for item in results]) if results else 0
        return text, confidence
        
    def extract_text_tesseract(self, image: np.ndarray) -> Tuple[str, float]:
        """Extract text using Tesseract"""
        # Convert to PIL Image
        pil_image = Image.fromarray(image)
        
        # Get text with confidence
        data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT)
        text = pytesseract.image_to_string(pil_image, lang='deu+eng')
        
        # Calculate confidence
        confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
        confidence = np.mean(confidences) if confidences else 0
        
        return text, confidence / 100
        
    def extract_tables(self, image: np.ndarray) -> List[pd.DataFrame]:
        """Extract tables from image using PaddleOCR"""
        tables = []
        
        if 'paddle' in self.engines:
            try:
                result = self.engines['paddle'].ocr(image, cls=True)
                # Process paddle results into tables
                # This is simplified - real implementation would use layout analysis
                if result:
                    # Group text by vertical position to form rows
                    pass
            except Exception as e:
                logger.error(f"Table extraction failed: {e}")
                
        return tables
        
    def process_image(self, image: np.ndarray) -> OCRResult:
        """Process a single image"""
        start_time = time.time()
        
        # Wait for resources
        self.resource_manager.wait_for_resources()
        
        # Preprocess image
        image = self._preprocess_image(image)
        
        # Try GPU OCR first, fallback to CPU
        text, confidence = "", 0.0
        
        try:
            if self.config.use_gpu and 'easyocr' in self.engines:
                text, confidence = self.extract_text_easyocr(image)
            else:
                raise Exception("Fallback to CPU")
        except Exception as e:
            logger.info(f"GPU OCR failed, using CPU: {e}")
            text, confidence = self.extract_text_tesseract(image)
            
        # Detect language
        language = self.detect_language(text)
        
        # Extract tables
        tables = self.extract_tables(image)
        
        # Convert to markdown
        markdown = self._to_markdown(text, tables)
        
        processing_time = time.time() - start_time
        
        return OCRResult(
            text=text,
            language=language,
            confidence=confidence,
            tables=tables,
            markdown=markdown,
            processing_time=processing_time,
            metadata={
                'engine': 'easyocr' if self.config.use_gpu else 'tesseract',
                'image_shape': image.shape
            }
        )
        
    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR"""
        # Convert to grayscale if needed
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
            
        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Threshold
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return thresh
        
    def _to_markdown(self, text: str, tables: List[pd.DataFrame]) -> str:
        """Convert text and tables to markdown"""
        markdown = text
        
        for i, table in enumerate(tables):
            markdown += f"\n\n### Table {i+1}\n\n"
            markdown += table.to_markdown(index=False)
            
        return markdown


class DocumentProcessor:
    """Process different document types"""
    
    def __init__(self, ocr_engine: OCREngine):
        self.ocr_engine = ocr_engine
        
    async def process_pdf(self, file_path: Path) -> List[OCRResult]:
        """Process PDF document"""
        results = []
        
        # Try native text extraction first
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            text = page.get_text()
            
            if text.strip():
                # Native text found
                results.append(OCRResult(
                    text=text,
                    language=self.ocr_engine.detect_language(text),
                    confidence=1.0,
                    tables=[],
                    markdown=text,
                    processing_time=0.1,
                    metadata={'page': page_num + 1, 'type': 'native'}
                ))
            else:
                # Convert to image and OCR
                pix = page.get_pixmap()
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                    pix.height, pix.width, pix.n
                )
                
                result = self.ocr_engine.process_image(img)
                result.metadata['page'] = page_num + 1
                results.append(result)
                
        doc.close()
        return results
        
    async def process_image(self, file_path: Path) -> OCRResult:
        """Process image file"""
        image = cv2.imread(str(file_path))
        return self.ocr_engine.process_image(image)
        
    async def process_docx(self, file_path: Path) -> OCRResult:
        """Process DOCX document"""
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        return OCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=[],
            markdown=text,
            processing_time=0.1,
            metadata={'type': 'docx', 'paragraphs': len(doc.paragraphs)}
        )
        
    async def process_txt(self, file_path: Path) -> OCRResult:
        """Process text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        return OCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=[],
            markdown=text,
            processing_time=0.01,
            metadata={'type': 'txt', 'size': len(text)}
        )


class OCRSystem:
    """Main OCR System with concurrent processing"""
    
    def __init__(self, config: Optional[OCRConfig] = None):
        self.config = config or OCRConfig()
        self.ocr_engine = OCREngine(self.config)
        self.document_processor = DocumentProcessor(self.ocr_engine)
        self.executor = ProcessPoolExecutor(max_workers=self.config.max_workers)
        
    async def process_file(self, file_path: str) -> OCRResult:
        """Process a single file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Determine file type and process
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            results = await self.document_processor.process_pdf(path)
            # Combine results from all pages
            combined_text = '\n\n'.join([r.text for r in results])
            combined_tables = []
            for r in results:
                combined_tables.extend(r.tables)
                
            return OCRResult(
                text=combined_text,
                language=results[0].language if results else 'en',
                confidence=np.mean([r.confidence for r in results]) if results else 0,
                tables=combined_tables,
                markdown='\n\n'.join([r.markdown for r in results]),
                processing_time=sum([r.processing_time for r in results]),
                metadata={'pages': len(results), 'type': 'pdf'}
            )
            
        elif suffix in ['.png', '.jpg', '.jpeg']:
            return await self.document_processor.process_image(path)
            
        elif suffix == '.docx':
            return await self.document_processor.process_docx(path)
            
        elif suffix == '.txt':
            return await self.document_processor.process_txt(path)
            
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
            
    async def process_batch(self, file_paths: List[str]) -> List[OCRResult]:
        """Process multiple files concurrently"""
        tasks = [self.process_file(fp) for fp in file_paths]
        return await asyncio.gather(*tasks)


def get_resource_info():
    """Get current resource usage information"""
    process = psutil.Process()
    return {
        'cpu_percent': process.cpu_percent(interval=1),
        'ram_mb': process.memory_info().rss / (1024 * 1024),
        'ram_gb': process.memory_info().rss / (1024 * 1024 * 1024),
        'ram_percent': process.memory_percent(),
        'threads': process.num_threads()
    }


def main():
    """Example usage"""
    print(f"OCR System Initialized")
    print(f"Resource Limits: CPU {MAX_CPU_PERCENT}%, RAM {MAX_RAM_PERCENT}% (Max {MAX_RAM_GB}GB)")
    print(f"GPU Available: {torch.cuda.is_available()}")
    
    # Show current resources
    resources = get_resource_info()
    print(f"\nCurrent Resources:")
    print(f"  CPU: {resources['cpu_percent']:.1f}%")
    print(f"  RAM: {resources['ram_gb']:.2f}GB ({resources['ram_percent']:.1f}%)")
    

if __name__ == "__main__":
    main()