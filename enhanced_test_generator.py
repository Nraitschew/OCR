#!/usr/bin/env python3
"""
Enhanced Test Document Generator with Complex Tables and Structures
"""

import os
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import numpy as np
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import cv2
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import logging

logger = logging.getLogger(__name__)

# Complex German texts with special formatting
COMPLEX_GERMAN_TEXTS = {
    "invoice": {
        "header": "RECHNUNG",
        "company": "Müller & Söhne GmbH",
        "address": "Königstraße 42, 80331 München",
        "items": [
            {"pos": 1, "artikel": "Bürostühle", "menge": 5, "preis": "125,50 €", "gesamt": "627,50 €"},
            {"pos": 2, "artikel": "Schreibtischlampen", "menge": 8, "preis": "45,99 €", "gesamt": "367,92 €"},
            {"pos": 3, "artikel": "Aktenordner (groß)", "menge": 20, "preis": "3,49 €", "gesamt": "69,80 €"},
            {"pos": 4, "artikel": "Druckerpapier A4", "menge": 10, "preis": "4,99 €", "gesamt": "49,90 €"},
        ],
        "subtotal": "1.115,12 €",
        "tax": "212,87 €",
        "total": "1.327,99 €"
    },
    "scientific": {
        "title": "Über die Wirkung von Säuren auf Metalle",
        "abstract": "Diese Untersuchung befasst sich mit der Korrosion verschiedener Metalle unter Einfluss von Säuren unterschiedlicher Konzentration.",
        "table_title": "Tabelle 1: Korrosionsraten verschiedener Metalle",
        "data": [
            ["Metall", "pH-Wert", "Temperatur (°C)", "Korrosionsrate (mm/Jahr)"],
            ["Eisen (Fe)", "2.5", "25", "0.127"],
            ["Kupfer (Cu)", "2.5", "25", "0.089"],
            ["Zink (Zn)", "2.5", "25", "0.234"],
            ["Aluminium (Al)", "2.5", "25", "0.045"],
        ]
    },
    "financial_report": {
        "title": "Finanzübersicht Q3 2024",
        "sections": {
            "revenue": "Umsatzerlöse",
            "expenses": "Aufwendungen",
            "profit": "Gewinn"
        },
        "quarters": ["Q1", "Q2", "Q3", "Q4"],
        "data": {
            "Umsatz": ["2.450.000 €", "2.890.000 €", "3.120.000 €", ""],
            "Kosten": ["1.980.000 €", "2.100.000 €", "2.340.000 €", ""],
            "Gewinn": ["470.000 €", "790.000 €", "780.000 €", ""],
        }
    }
}


class EnhancedTestDocumentGenerator:
    """Generate complex test documents with tables and structures"""
    
    def __init__(self, output_dir: str = "test_documents_enhanced"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def generate_all(self):
        """Generate all enhanced test documents"""
        logger.info("Generating enhanced test documents...")
        
        # Generate complex tables
        self.generate_complex_tables()
        
        # Generate invoice documents
        self.generate_invoice_documents()
        
        # Generate scientific documents
        self.generate_scientific_documents()
        
        # Generate financial reports
        self.generate_financial_reports()
        
        # Generate mixed content documents
        self.generate_mixed_content_documents()
        
        # Generate nested table documents
        self.generate_nested_tables()
        
        # Generate formatted documents
        self.generate_formatted_documents()
        
        logger.info(f"Enhanced test documents generated in {self.output_dir}")
        
    def generate_complex_tables(self):
        """Generate documents with complex table structures"""
        
        # Multi-column table with merged cells
        self._create_multi_column_table_pdf()
        self._create_multi_column_table_docx()
        self._create_multi_column_table_image()
        
    def _create_multi_column_table_pdf(self):
        """Create PDF with complex multi-column table"""
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', size=14)
        
        # Title
        pdf.cell(0, 10, 'Komplexe Tabelle - Verkaufsübersicht', ln=True, align='C')
        pdf.ln(5)
        
        # Complex table header
        pdf.set_font('DejaVu', size=10)
        pdf.set_fill_color(200, 200, 200)
        
        # Header row 1
        pdf.cell(40, 10, 'Produkt', border=1, fill=True)
        pdf.cell(60, 5, 'Verkaufszahlen 2024', border=1, fill=True, align='C')
        pdf.cell(60, 5, 'Umsatz in €', border=1, fill=True, align='C')
        pdf.ln(5)
        
        # Header row 2
        pdf.cell(40, 5, '', border=0)
        pdf.cell(20, 5, 'Q1', border=1, fill=True, align='C')
        pdf.cell(20, 5, 'Q2', border=1, fill=True, align='C')
        pdf.cell(20, 5, 'Q3', border=1, fill=True, align='C')
        pdf.cell(20, 5, 'Q1', border=1, fill=True, align='C')
        pdf.cell(20, 5, 'Q2', border=1, fill=True, align='C')
        pdf.cell(20, 5, 'Q3', border=1, fill=True, align='C')
        pdf.ln(5)
        
        # Data rows
        products = [
            ("Bücher", "1.250", "1.480", "1.320", "31.250", "37.000", "33.000"),
            ("Hörbücher", "890", "920", "1.050", "17.800", "18.400", "21.000"),
            ("E-Books", "2.100", "2.450", "2.890", "21.000", "24.500", "28.900"),
            ("Zeitschriften", "450", "380", "410", "4.500", "3.800", "4.100"),
        ]
        
        for product in products:
            pdf.cell(40, 8, product[0], border=1)
            for value in product[1:]:
                pdf.cell(20, 8, value, border=1, align='R')
            pdf.ln()
            
        # Summary row
        pdf.set_fill_color(220, 220, 220)
        pdf.cell(40, 8, 'Gesamt', border=1, fill=True)
        pdf.cell(20, 8, '4.690', border=1, fill=True, align='R')
        pdf.cell(20, 8, '5.230', border=1, fill=True, align='R')
        pdf.cell(20, 8, '5.670', border=1, fill=True, align='R')
        pdf.cell(20, 8, '74.550', border=1, fill=True, align='R')
        pdf.cell(20, 8, '83.700', border=1, fill=True, align='R')
        pdf.cell(20, 8, '87.000', border=1, fill=True, align='R')
        
        pdf.output(str(self.output_dir / "complex_table_multicolumn.pdf"))
        
    def _create_multi_column_table_docx(self):
        """Create DOCX with complex multi-column table"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Komplexe Tabelle - Verkaufsübersicht', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create table with merged cells
        table = doc.add_table(rows=7, cols=7)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Merge cells for headers
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(0, 3))
        table.cell(0, 4).merge(table.cell(0, 6))
        
        # Fill headers
        table.cell(0, 0).text = 'Produkt'
        table.cell(0, 1).text = 'Verkaufszahlen 2024'
        table.cell(0, 4).text = 'Umsatz in €'
        
        # Sub-headers
        headers = ['Q1', 'Q2', 'Q3']
        for i, header in enumerate(headers):
            table.cell(1, i+1).text = header
            table.cell(1, i+4).text = header
            
        # Data
        products = [
            ("Bücher", "1.250", "1.480", "1.320", "31.250", "37.000", "33.000"),
            ("Hörbücher", "890", "920", "1.050", "17.800", "18.400", "21.000"),
            ("E-Books", "2.100", "2.450", "2.890", "21.000", "24.500", "28.900"),
            ("Zeitschriften", "450", "380", "410", "4.500", "3.800", "4.100"),
            ("Gesamt", "4.690", "5.230", "5.670", "74.550", "83.700", "87.000"),
        ]
        
        for row_idx, product in enumerate(products):
            for col_idx, value in enumerate(product):
                table.cell(row_idx + 2, col_idx).text = value
                
        doc.save(str(self.output_dir / "complex_table_multicolumn.docx"))
        
    def _create_multi_column_table_image(self):
        """Create image with complex table"""
        # Create figure
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.axis('tight')
        ax.axis('off')
        
        # Create table data
        headers = [['Produkt', 'Verkaufszahlen 2024', '', '', 'Umsatz in €', '', ''],
                   ['', 'Q1', 'Q2', 'Q3', 'Q1', 'Q2', 'Q3']]
        
        data = [
            ["Bücher", "1.250", "1.480", "1.320", "31.250 €", "37.000 €", "33.000 €"],
            ["Hörbücher", "890", "920", "1.050", "17.800 €", "18.400 €", "21.000 €"],
            ["E-Books", "2.100", "2.450", "2.890", "21.000 €", "24.500 €", "28.900 €"],
            ["Zeitschriften", "450", "380", "410", "4.500 €", "3.800 €", "4.100 €"],
            ["Gesamt", "4.690", "5.230", "5.670", "74.550 €", "83.700 €", "87.000 €"],
        ]
        
        # Create table
        table = ax.table(cellText=headers + data,
                        cellLoc='center',
                        loc='center',
                        colWidths=[0.15, 0.12, 0.12, 0.12, 0.15, 0.15, 0.15])
        
        table.auto_set_font_size(False)
        table.set_fontsize(10)
        table.scale(1, 2)
        
        # Style headers
        for i in range(7):
            table[(0, i)].set_facecolor('#CCCCCC')
            table[(1, i)].set_facecolor('#DDDDDD')
            
        # Style total row
        for i in range(7):
            table[(6, i)].set_facecolor('#EEEEEE')
            
        plt.title('Komplexe Tabelle - Verkaufsübersicht', fontsize=16, pad=20)
        plt.savefig(str(self.output_dir / "complex_table_multicolumn.png"), 
                    dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        
    def generate_invoice_documents(self):
        """Generate invoice-style documents"""
        invoice_data = COMPLEX_GERMAN_TEXTS["invoice"]
        
        # PDF Invoice
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
        pdf.add_font('DejaVu', 'B', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', uni=True)
        
        # Header
        pdf.set_font('DejaVu', 'B', size=20)
        pdf.cell(0, 15, invoice_data["header"], ln=True, align='C')
        
        # Company info
        pdf.set_font('DejaVu', size=12)
        pdf.cell(0, 8, invoice_data["company"], ln=True)
        pdf.cell(0, 8, invoice_data["address"], ln=True)
        pdf.ln(10)
        
        # Invoice details
        pdf.set_font('DejaVu', size=10)
        pdf.cell(30, 8, 'Rechnungsnr:', border=0)
        pdf.cell(50, 8, 'RE-2024-001234', border=0)
        pdf.cell(30, 8, 'Datum:', border=0)
        pdf.cell(30, 8, '15.11.2024', border=0)
        pdf.ln(10)
        
        # Items table
        pdf.set_font('DejaVu', 'B', size=10)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(15, 10, 'Pos', border=1, fill=True, align='C')
        pdf.cell(80, 10, 'Artikel', border=1, fill=True)
        pdf.cell(25, 10, 'Menge', border=1, fill=True, align='C')
        pdf.cell(35, 10, 'Einzelpreis', border=1, fill=True, align='R')
        pdf.cell(35, 10, 'Gesamt', border=1, fill=True, align='R')
        pdf.ln()
        
        # Items
        pdf.set_font('DejaVu', size=10)
        for item in invoice_data["items"]:
            pdf.cell(15, 8, str(item["pos"]), border=1, align='C')
            pdf.cell(80, 8, item["artikel"], border=1)
            pdf.cell(25, 8, str(item["menge"]), border=1, align='C')
            pdf.cell(35, 8, item["preis"], border=1, align='R')
            pdf.cell(35, 8, item["gesamt"], border=1, align='R')
            pdf.ln()
            
        # Summary
        pdf.ln(5)
        pdf.cell(120, 8, '', border=0)
        pdf.cell(35, 8, 'Zwischensumme:', border=0)
        pdf.cell(35, 8, invoice_data["subtotal"], border=0, align='R')
        pdf.ln()
        pdf.cell(120, 8, '', border=0)
        pdf.cell(35, 8, 'MwSt (19%):', border=0)
        pdf.cell(35, 8, invoice_data["tax"], border=0, align='R')
        pdf.ln()
        pdf.set_font('DejaVu', 'B', size=12)
        pdf.cell(120, 10, '', border=0)
        pdf.cell(35, 10, 'Gesamtbetrag:', border='T')
        pdf.cell(35, 10, invoice_data["total"], border='T', align='R')
        
        pdf.output(str(self.output_dir / "invoice_complex.pdf"))
        
    def generate_scientific_documents(self):
        """Generate scientific documents with tables and formulas"""
        sci_data = COMPLEX_GERMAN_TEXTS["scientific"]
        
        # Create scientific PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
        pdf.add_font('DejaVu', 'B', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', uni=True)
        
        # Title
        pdf.set_font('DejaVu', 'B', size=16)
        pdf.multi_cell(0, 10, sci_data["title"], align='C')
        pdf.ln(5)
        
        # Abstract
        pdf.set_font('DejaVu', 'B', size=12)
        pdf.cell(0, 10, 'Zusammenfassung', ln=True)
        pdf.set_font('DejaVu', size=10)
        pdf.multi_cell(0, 5, sci_data["abstract"])
        pdf.ln(10)
        
        # Table
        pdf.set_font('DejaVu', 'B', size=10)
        pdf.cell(0, 8, sci_data["table_title"], ln=True)
        pdf.ln(3)
        
        # Table data
        pdf.set_font('DejaVu', size=9)
        col_widths = [40, 30, 40, 50]
        
        # Header
        pdf.set_fill_color(200, 200, 200)
        for i, header in enumerate(sci_data["data"][0]):
            pdf.cell(col_widths[i], 8, header, border=1, fill=True, align='C')
        pdf.ln()
        
        # Data rows
        for row in sci_data["data"][1:]:
            for i, cell in enumerate(row):
                pdf.cell(col_widths[i], 8, cell, border=1, align='C')
            pdf.ln()
            
        # Add chemical formulas
        pdf.ln(10)
        pdf.set_font('DejaVu', size=10)
        pdf.cell(0, 8, 'Chemische Reaktionen:', ln=True)
        pdf.cell(0, 8, 'Fe + 2HCl → FeCl₂ + H₂', ln=True)
        pdf.cell(0, 8, 'Zn + H₂SO₄ → ZnSO₄ + H₂', ln=True)
        
        pdf.output(str(self.output_dir / "scientific_document.pdf"))
        
    def generate_financial_reports(self):
        """Generate financial report documents"""
        fin_data = COMPLEX_GERMAN_TEXTS["financial_report"]
        
        # Create DOCX financial report
        doc = Document()
        
        # Title
        title = doc.add_heading(fin_data["title"], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create financial table
        table = doc.add_table(rows=4, cols=5)
        table.style = 'Medium Grid 3 Accent 1'
        
        # Headers
        headers = ['Kategorie'] + fin_data["quarters"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            
        # Data
        row_idx = 1
        for category, values in fin_data["data"].items():
            table.cell(row_idx, 0).text = category
            for col_idx, value in enumerate(values):
                table.cell(row_idx, col_idx + 1).text = value
            row_idx += 1
            
        # Add chart placeholder
        doc.add_paragraph()
        doc.add_paragraph('Grafische Darstellung der Quartalsübersicht:')
        
        # Create a simple bar chart as image
        self._create_financial_chart()
        doc.add_picture(str(self.output_dir / "financial_chart.png"), width=Inches(5))
        
        doc.save(str(self.output_dir / "financial_report.docx"))
        
    def _create_financial_chart(self):
        """Create financial chart"""
        quarters = ['Q1', 'Q2', 'Q3']
        revenue = [2450, 2890, 3120]
        costs = [1980, 2100, 2340]
        profit = [470, 790, 780]
        
        x = np.arange(len(quarters))
        width = 0.25
        
        fig, ax = plt.subplots(figsize=(10, 6))
        bars1 = ax.bar(x - width, revenue, width, label='Umsatz')
        bars2 = ax.bar(x, costs, width, label='Kosten')
        bars3 = ax.bar(x + width, profit, width, label='Gewinn')
        
        ax.set_xlabel('Quartal')
        ax.set_ylabel('Betrag (in Tausend €)')
        ax.set_title('Finanzübersicht Q1-Q3 2024')
        ax.set_xticks(x)
        ax.set_xticklabels(quarters)
        ax.legend()
        
        # Add value labels on bars
        def autolabel(bars):
            for bar in bars:
                height = bar.get_height()
                ax.annotate(f'{height}',
                           xy=(bar.get_x() + bar.get_width() / 2, height),
                           xytext=(0, 3),
                           textcoords="offset points",
                           ha='center', va='bottom',
                           fontsize=8)
        
        autolabel(bars1)
        autolabel(bars2)
        autolabel(bars3)
        
        plt.tight_layout()
        plt.savefig(str(self.output_dir / "financial_chart.png"), dpi=150, facecolor='white')
        plt.close()
        
    def generate_mixed_content_documents(self):
        """Generate documents with mixed content (text, tables, images)"""
        # Create a complex mixed document
        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
        
        # Page header
        pdf.set_font('DejaVu', size=8)
        pdf.cell(0, 5, 'Jahresbericht 2024 - Seite 1', align='R')
        pdf.ln(10)
        
        # Main title
        pdf.set_font('DejaVu', size=18)
        pdf.cell(0, 10, 'Geschäftsbericht der Müller GmbH', ln=True, align='C')
        pdf.ln(5)
        
        # Section 1: Text
        pdf.set_font('DejaVu', size=12)
        pdf.cell(0, 8, '1. Überblick', ln=True)
        pdf.set_font('DejaVu', size=10)
        pdf.multi_cell(0, 5, 
            'Das Geschäftsjahr 2024 war für die Müller GmbH ein äußerst erfolgreiches Jahr. '
            'Mit einer Umsatzsteigerung von 23% konnten wir unsere Marktposition deutlich ausbauen. '
            'Besonders erfreulich war die Entwicklung im Bereich der digitalen Dienstleistungen.')
        pdf.ln(5)
        
        # Section 2: Table
        pdf.set_font('DejaVu', size=12)
        pdf.cell(0, 8, '2. Kennzahlen im Überblick', ln=True)
        pdf.ln(3)
        
        # Kennzahlen table
        pdf.set_font('DejaVu', size=9)
        pdf.set_fill_color(240, 240, 240)
        
        kennzahlen = [
            ['Kennzahl', '2023', '2024', 'Veränderung'],
            ['Umsatz (Mio. €)', '12,5', '15,4', '+23%'],
            ['Mitarbeiter', '89', '112', '+26%'],
            ['Standorte', '3', '4', '+33%'],
            ['Kundenzufriedenheit', '4,2/5', '4,6/5', '+9,5%']
        ]
        
        for i, row in enumerate(kennzahlen):
            for j, cell in enumerate(row):
                if i == 0:
                    pdf.cell(45, 8, cell, border=1, fill=True, align='C')
                else:
                    pdf.cell(45, 8, cell, border=1, align='C')
            pdf.ln()
            
        # Add page break for more content
        pdf.add_page()
        
        # Page header
        pdf.set_font('DejaVu', size=8)
        pdf.cell(0, 5, 'Jahresbericht 2024 - Seite 2', align='R')
        pdf.ln(10)
        
        # Section 3: Regional analysis
        pdf.set_font('DejaVu', size=12)
        pdf.cell(0, 8, '3. Regionale Verteilung', ln=True)
        pdf.ln(3)
        
        # Regional table with special characters
        regional = [
            ['Region', 'Städte', 'Umsatzanteil'],
            ['Süddeutschland', 'München, Nürnberg', '45%'],
            ['Norddeutschland', 'Hamburg, Lübeck', '25%'],
            ['Westdeutschland', 'Köln, Düsseldorf', '20%'],
            ['Ostdeutschland', 'Leipzig, Görlitz', '10%']
        ]
        
        pdf.set_font('DejaVu', size=9)
        col_widths = [50, 70, 40]
        
        for i, row in enumerate(regional):
            for j, cell in enumerate(row):
                if i == 0:
                    pdf.cell(col_widths[j], 8, cell, border=1, fill=True, align='C')
                else:
                    pdf.cell(col_widths[j], 8, cell, border=1)
            pdf.ln()
            
        pdf.output(str(self.output_dir / "mixed_content_report.pdf"))
        
    def generate_nested_tables(self):
        """Generate documents with nested table structures"""
        # Create image with nested tables visualization
        fig, ax = plt.subplots(figsize=(14, 10))
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.axis('off')
        
        # Main title
        ax.text(5, 9.5, 'Organisationsstruktur - Müller & Partner AG', 
                fontsize=16, ha='center', weight='bold')
        
        # Main table structure
        main_rect = patches.Rectangle((0.5, 6), 9, 3, 
                                     linewidth=2, edgecolor='black', 
                                     facecolor='lightblue', alpha=0.3)
        ax.add_patch(main_rect)
        ax.text(5, 7.5, 'Geschäftsführung', fontsize=12, ha='center', weight='bold')
        
        # Nested divisions
        divisions = [
            ('Vertrieb', 1, 4, 'lightgreen'),
            ('Produktion', 3.5, 4, 'lightyellow'),
            ('Verwaltung', 6, 4, 'lightcoral'),
            ('Forschung', 8.5, 4, 'lightgray')
        ]
        
        for div_name, x, y, color in divisions:
            rect = patches.Rectangle((x-0.75, y-0.5), 1.5, 1.5,
                                    linewidth=1, edgecolor='black',
                                    facecolor=color, alpha=0.5)
            ax.add_patch(rect)
            ax.text(x, y+0.5, div_name, fontsize=10, ha='center', weight='bold')
            
            # Sub-departments
            if div_name == 'Vertrieb':
                subdepts = ['Inland', 'Ausland']
                for i, subdept in enumerate(subdepts):
                    subrect = patches.Rectangle((x-0.6+i*0.6, y-0.3), 0.5, 0.6,
                                              linewidth=0.5, edgecolor='gray',
                                              facecolor='white', alpha=0.8)
                    ax.add_patch(subrect)
                    ax.text(x-0.35+i*0.6, y, subdept, fontsize=8, ha='center')
                    
        # Add employee counts table
        ax.text(5, 2.5, 'Mitarbeiterverteilung', fontsize=12, ha='center', weight='bold')
        
        # Employee table
        emp_data = [
            ['Abteilung', 'Vollzeit', 'Teilzeit', 'Gesamt'],
            ['Vertrieb', '45', '12', '57'],
            ['Produktion', '78', '8', '86'],
            ['Verwaltung', '23', '15', '38'],
            ['Forschung', '34', '5', '39'],
            ['Gesamt', '180', '40', '220']
        ]
        
        table_y = 1.5
        cell_height = 0.3
        cell_widths = [2, 1.5, 1.5, 1.5]
        
        for i, row in enumerate(emp_data):
            for j, cell in enumerate(row):
                x_pos = 2 + sum(cell_widths[:j])
                y_pos = table_y - i * cell_height
                
                # Cell background
                if i == 0 or i == len(emp_data)-1:
                    cell_color = 'lightgray'
                else:
                    cell_color = 'white'
                    
                rect = patches.Rectangle((x_pos, y_pos), cell_widths[j], cell_height,
                                       linewidth=0.5, edgecolor='black',
                                       facecolor=cell_color, alpha=0.8)
                ax.add_patch(rect)
                
                # Cell text
                ax.text(x_pos + cell_widths[j]/2, y_pos + cell_height/2, cell,
                       fontsize=8, ha='center', va='center',
                       weight='bold' if i == 0 else 'normal')
                
        plt.savefig(str(self.output_dir / "nested_tables_org_chart.png"),
                    dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        
    def generate_formatted_documents(self):
        """Generate documents with complex formatting"""
        # Create a formatted price list
        doc = Document()
        
        # Header with formatting
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run('PREISLISTE 2024')
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Gültig ab 01.01.2024')
        run.font.size = Pt(12)
        run.font.italic = True
        
        doc.add_paragraph()
        
        # Category 1
        cat1 = doc.add_paragraph()
        run = cat1.add_run('Kategorie A - Bürobedarf')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.underline = True
        
        # Price table with formatting
        table1 = doc.add_table(rows=5, cols=4)
        table1.style = 'Light List Accent 1'
        
        # Headers
        headers = ['Art.-Nr.', 'Bezeichnung', 'Einheit', 'Preis (€)']
        header_cells = table1.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make headers bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        # Products with special characters
        products = [
            ['A-001', 'Büroklammern (groß)', '100 Stk.', '2,49'],
            ['A-002', 'Heftgerät (25 Blatt)', 'Stück', '12,99'],
            ['A-003', 'Locher (für 30 Blatt)', 'Stück', '8,49'],
            ['A-004', 'Textmarker (4er-Set)', 'Set', '4,99']
        ]
        
        for i, product in enumerate(products):
            row = table1.rows[i + 1]
            for j, value in enumerate(product):
                row.cells[j].text = value
                if j == 3:  # Price column
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
        doc.add_paragraph()
        
        # Category 2 with different formatting
        cat2 = doc.add_paragraph()
        run = cat2.add_run('Kategorie B - Büromöbel')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.underline = True
        
        # Second table with merged cells
        table2 = doc.add_table(rows=6, cols=5)
        table2.style = 'Medium Grid 3 Accent 2'
        
        # Merge cells for special offers
        table2.cell(0, 0).merge(table2.cell(0, 4))
        table2.cell(0, 0).text = 'SONDERANGEBOTE - Nur für kurze Zeit!'
        
        # Make special offer text red and bold
        for paragraph in table2.cell(0, 0).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                
        # Headers for furniture
        headers2 = ['Art.-Nr.', 'Möbelstück', 'Maße (BxTxH)', 'Regulär', 'Angebot']
        for i, header in enumerate(headers2):
            table2.cell(1, i).text = header
            
        # Furniture items
        furniture = [
            ['B-101', 'Schreibtisch "Ergonomic"', '160x80x75 cm', '349,00 €', '279,00 €'],
            ['B-102', 'Bürostuhl "Komfort Plus"', '65x65x120 cm', '199,00 €', '159,00 €'],
            ['B-103', 'Aktenschrank (3 Fächer)', '80x40x110 cm', '149,00 €', '119,00 €'],
            ['B-104', 'Besprechungstisch (6 Pers.)', '200x100x75 cm', '499,00 €', '399,00 €']
        ]
        
        for i, item in enumerate(furniture):
            row = table2.rows[i + 2]
            for j, value in enumerate(item):
                row.cells[j].text = value
                if j >= 3:  # Price columns
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if j == 4:  # Sale price in red
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.bold = True
                            
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('* Alle Preise verstehen sich zzgl. 19% MwSt.')
        run.font.size = Pt(10)
        run.font.italic = True
        
        doc.save(str(self.output_dir / "formatted_price_list.docx"))
        
        
def main():
    """Generate all enhanced test documents"""
    generator = EnhancedTestDocumentGenerator()
    generator.generate_all()
    
    # List generated files
    print("\nGenerated enhanced test documents:")
    for file in sorted(generator.output_dir.iterdir()):
        if file.is_file():
            size_kb = file.stat().st_size / 1024
            print(f"  - {file.name} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()