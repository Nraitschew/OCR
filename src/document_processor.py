"""
Document processors for different file formats
"""
import os
import io
import logging
from typing import Dict, List, Optional, Union
from pathlib import Path
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
import numpy as np

from .config import (
    SUPPORTED_FORMATS, IMAGE_DPI, IMAGE_MAX_SIZE,
    TABLE_DETECTION_ENABLED, BATCH_SIZE
)
from .ocr_engine import OCREngine
from .resource_manager import ResourceContext

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Base class for document processing"""
    
    def __init__(self, ocr_engine: OCREngine):
        self.ocr_engine = ocr_engine
        self.resource_manager = ocr_engine.resource_manager
        
    def process_file(self, file_path: str) -> Dict:
        """Process a file and return OCR results"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = file_path.suffix.lower()
        
        if ext not in SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {ext}")
            
        logger.info(f"Processing {ext} file: {file_path}")
        
        # Route to appropriate processor
        if ext == '.pdf':
            return self._process_pdf(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self._process_image(file_path)
        elif ext == '.docx':
            return self._process_docx(file_path)
        elif ext == '.txt':
            return self._process_text(file_path)
        else:
            raise ValueError(f"No processor for format: {ext}")
            
    def _process_pdf(self, file_path: Path) -> Dict:
        """Process PDF file (both native and scanned)"""
        results = {
            'file_path': str(file_path),
            'file_type': 'pdf',
            'pages': [],
            'total_pages': 0,
            'has_text': False,
            'is_scanned': False
        }
        
        with ResourceContext(self.resource_manager, memory_mb=500):
            # Open PDF
            pdf_document = fitz.open(str(file_path))
            results['total_pages'] = len(pdf_document)
            
            # Check if PDF has extractable text
            for page_num in range(min(3, len(pdf_document))):  # Check first 3 pages
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    results['has_text'] = True
                    break
                    
            # Process each page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_result = {
                    'page_number': page_num + 1,
                    'text': '',
                    'tables': [],
                    'has_native_text': False,
                    'ocr_applied': False
                }
                
                # Try to extract native text first
                native_text = page.get_text()
                
                if native_text.strip():
                    page_result['text'] = native_text
                    page_result['has_native_text'] = True
                    logger.debug(f"Extracted native text from page {page_num + 1}")
                else:
                    # No native text, convert to image and OCR
                    results['is_scanned'] = True
                    
                    # Convert page to image
                    pix = page.get_pixmap(dpi=IMAGE_DPI)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Apply OCR
                    ocr_result = self.ocr_engine.process_image(
                        img,
                        detect_tables=TABLE_DETECTION_ENABLED
                    )
                    
                    page_result['text'] = ocr_result['text']
                    page_result['tables'] = ocr_result.get('tables', [])
                    page_result['ocr_applied'] = True
                    page_result['ocr_confidence'] = ocr_result.get('confidence', 0)
                    page_result['ocr_engine'] = ocr_result.get('engine_used', '')
                    
                    logger.debug(f"Applied OCR to page {page_num + 1}")
                    
                results['pages'].append(page_result)
                
            pdf_document.close()
            
        return results
        
    def _process_image(self, file_path: Path) -> Dict:
        """Process image file"""
        results = {
            'file_path': str(file_path),
            'file_type': 'image',
            'format': file_path.suffix.lower()[1:]  # Remove dot
        }
        
        with ResourceContext(self.resource_manager, memory_mb=300):
            # Open and potentially resize image
            img = Image.open(str(file_path))
            
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
                
            # Resize if too large
            if img.size[0] > IMAGE_MAX_SIZE[0] or img.size[1] > IMAGE_MAX_SIZE[1]:
                img.thumbnail(IMAGE_MAX_SIZE, Image.Resampling.LANCZOS)
                logger.debug(f"Resized image to {img.size}")
                
            # Process with OCR
            ocr_result = self.ocr_engine.process_image(
                img,
                detect_tables=TABLE_DETECTION_ENABLED
            )
            
            results.update(ocr_result)
            
        return results
        
    def _process_docx(self, file_path: Path) -> Dict:
        """Process DOCX file"""
        results = {
            'file_path': str(file_path),
            'file_type': 'docx',
            'paragraphs': [],
            'tables': []
        }
        
        with ResourceContext(self.resource_manager, memory_mb=200):
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    results['paragraphs'].append(para.text)
                    
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                    
                if table_data:
                    results['tables'].append({
                        'data': table_data,
                        'confidence': 1.0  # Native extraction
                    })
                    
            # Combine all text
            results['text'] = '\n\n'.join(results['paragraphs'])
            
        return results
        
    def _process_text(self, file_path: Path) -> Dict:
        """Process text file"""
        results = {
            'file_path': str(file_path),
            'file_type': 'text',
            'encoding': 'utf-8'
        }
        
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    results['text'] = f.read()
                    results['encoding'] = encoding
                    break
            except UnicodeDecodeError:
                continue
        else:
            # If all encodings fail, read as binary and try to decode
            with open(file_path, 'rb') as f:
                content = f.read()
                results['text'] = content.decode('utf-8', errors='replace')
                results['encoding'] = 'utf-8 (with errors)'
                
        return results
        
    def process_batch(self, file_paths: List[str]) -> List[Dict]:
        """Process multiple files in batch"""
        results = []
        
        # Process in batches to manage memory
        for i in range(0, len(file_paths), BATCH_SIZE):
            batch = file_paths[i:i + BATCH_SIZE]
            
            for file_path in batch:
                try:
                    result = self.process_file(file_path)
                    results.append(result)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    results.append({
                        'file_path': file_path,
                        'error': str(e)
                    })
                    
        return results


class PDFProcessor:
    """Specialized processor for PDF files with advanced features"""
    
    @staticmethod
    def convert_pdf_to_images(pdf_path: str, dpi: int = IMAGE_DPI) -> List[Image.Image]:
        """Convert PDF pages to images"""
        try:
            images = convert_from_path(pdf_path, dpi=dpi)
            return images
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            # Fallback to PyMuPDF
            return PDFProcessor._convert_with_pymupdf(pdf_path, dpi)
            
    @staticmethod
    def _convert_with_pymupdf(pdf_path: str, dpi: int) -> List[Image.Image]:
        """Convert PDF to images using PyMuPDF as fallback"""
        images = []
        pdf_document = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            pix = page.get_pixmap(dpi=dpi)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            images.append(img)
            
        pdf_document.close()
        return images