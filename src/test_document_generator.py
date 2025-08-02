"""
Generate and download test documents for OCR testing
"""
import os
import requests
from PIL import Image, ImageDraw, ImageFont
import numpy as np
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches
import logging
from .config import TEST_DOCUMENTS_DIR

logger = logging.getLogger(__name__)


class TestDocumentGenerator:
    """Generate test documents with German text and various formats"""
    
    def __init__(self):
        self.output_dir = TEST_DOCUMENTS_DIR
        os.makedirs(self.output_dir, exist_ok=True)
        
        # German test texts with umlauts
        self.german_texts = [
            "Übung macht den Meister. Äpfel und Öfen sind schön.",
            "Die Größe der Bäume überrascht die Gäste im Frühjahr.",
            "Müller, Schäfer und Jäger treffen sich am Flüßchen.",
            "Das Mädchen trägt ein hübsches Kleid zur Hochzeit.",
            "Straßenbahn fährt durch München, Köln und Düsseldorf."
        ]
        
        self.table_data = [
            ["Stadt", "Einwohner", "Fläche (km²)", "Bundesland"],
            ["München", "1.471.508", "310,74", "Bayern"],
            ["Köln", "1.085.664", "405,02", "Nordrhein-Westfalen"],
            ["Düsseldorf", "619.294", "217,41", "Nordrhein-Westfalen"],
            ["Nürnberg", "518.365", "186,38", "Bayern"],
            ["Würzburg", "127.880", "87,63", "Bayern"]
        ]
        
    def generate_all_test_documents(self):
        """Generate all test documents"""
        logger.info("Generating test documents...")
        
        # Generate different format documents
        self._generate_text_file()
        self._generate_pdf_native()
        self._generate_pdf_scanned()
        self._generate_images()
        self._generate_docx()
        self._download_sample_documents()
        
        logger.info("Test documents generated successfully")
        
    def _generate_text_file(self):
        """Generate plain text file with German text"""
        filepath = os.path.join(self.output_dir, "german_text.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("Deutsche Textdatei mit Umlauten\n")
            f.write("=" * 40 + "\n\n")
            for text in self.german_texts:
                f.write(text + "\n\n")
            f.write("\nTabelle der größten Städte:\n")
            f.write("-" * 40 + "\n")
            for row in self.table_data:
                f.write(" | ".join(row) + "\n")
        logger.info(f"Created text file: {filepath}")
        
    def _generate_pdf_native(self):
        """Generate native PDF with German text and table"""
        filepath = os.path.join(self.output_dir, "german_native.pdf")
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title = Paragraph("Deutsche PDF-Datei mit Umlauten", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Add German text
        for text in self.german_texts:
            para = Paragraph(text, styles['Normal'])
            story.append(para)
            story.append(Spacer(1, 6))
            
        # Add table
        story.append(Spacer(1, 12))
        story.append(Paragraph("Größte Städte Deutschlands:", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        table = Table(self.table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
        
        doc.build(story)
        logger.info(f"Created native PDF: {filepath}")
        
    def _generate_pdf_scanned(self):
        """Generate scanned-style PDF (image-based)"""
        # First create an image with text
        img_path = os.path.join(self.output_dir, "temp_scan.png")
        img = Image.new('RGB', (2480, 3508), color='white')  # A4 at 300 DPI
        draw = ImageDraw.Draw(img)
        
        # Try to use a font, fallback to default if not available
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", 60)
            title_font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf", 80)
        except:
            font = ImageFont.load_default()
            title_font = font
            
        # Add some noise and rotation to simulate scan
        y_position = 200
        
        # Title
        draw.text((200, y_position), "Gescanntes Dokument mit Umlauten", fill='black', font=title_font)
        y_position += 150
        
        # Add German text
        for text in self.german_texts:
            draw.text((200, y_position), text, fill='black', font=font)
            y_position += 100
            
        # Add table
        y_position += 100
        draw.text((200, y_position), "Städtetabelle:", fill='black', font=title_font)
        y_position += 100
        
        for row in self.table_data:
            x_position = 200
            for cell in row:
                draw.text((x_position, y_position), cell, fill='black', font=font)
                x_position += 400
            y_position += 80
            
        # Add some noise to simulate scan artifacts
        pixels = np.array(img)
        noise = np.random.normal(0, 5, pixels.shape)
        pixels = np.clip(pixels + noise, 0, 255).astype(np.uint8)
        img = Image.fromarray(pixels)
        
        # Save as image first
        img.save(img_path, 'PNG')
        
        # Convert to PDF
        pdf_path = os.path.join(self.output_dir, "german_scanned.pdf")
        img.save(pdf_path, 'PDF')
        
        # Clean up temp image
        os.remove(img_path)
        
        logger.info(f"Created scanned PDF: {pdf_path}")
        
    def _generate_images(self):
        """Generate test images with German text"""
        for i, text in enumerate(self.german_texts[:3]):
            img = Image.new('RGB', (800, 200), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", 40)
            except:
                font = ImageFont.load_default()
                
            draw.text((50, 50), text, fill='black', font=font)
            
            # Save as both PNG and JPEG
            png_path = os.path.join(self.output_dir, f"german_text_{i+1}.png")
            jpg_path = os.path.join(self.output_dir, f"german_text_{i+1}.jpg")
            
            img.save(png_path, 'PNG')
            img.save(jpg_path, 'JPEG')
            
            logger.info(f"Created images: {png_path}, {jpg_path}")
            
    def _generate_docx(self):
        """Generate DOCX file with German text and table"""
        doc = Document()
        
        # Add title
        doc.add_heading('Deutsche Word-Datei mit Umlauten', 0)
        
        # Add German paragraphs
        for text in self.german_texts:
            doc.add_paragraph(text)
            
        # Add table
        doc.add_heading('Größte Städte Deutschlands', level=1)
        
        table = doc.add_table(rows=len(self.table_data), cols=len(self.table_data[0]))
        table.style = 'Light List Accent 1'
        
        for i, row_data in enumerate(self.table_data):
            for j, cell_data in enumerate(row_data):
                table.rows[i].cells[j].text = cell_data
                
        filepath = os.path.join(self.output_dir, "german_document.docx")
        doc.save(filepath)
        logger.info(f"Created DOCX file: {filepath}")
        
    def _download_sample_documents(self):
        """Download sample documents from the internet"""
        samples = [
            {
                'url': 'https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf',
                'filename': 'sample_english.pdf'
            },
            {
                'url': 'https://sample-videos.com/pdf/Sample-pdf-5mb.pdf',
                'filename': 'sample_large.pdf'
            }
        ]
        
        for sample in samples:
            try:
                filepath = os.path.join(self.output_dir, sample['filename'])
                response = requests.get(sample['url'], timeout=30)
                if response.status_code == 200:
                    with open(filepath, 'wb') as f:
                        f.write(response.content)
                    logger.info(f"Downloaded: {filepath}")
                else:
                    logger.warning(f"Failed to download: {sample['url']}")
            except Exception as e:
                logger.error(f"Error downloading {sample['url']}: {e}")
                
                
if __name__ == "__main__":
    generator = TestDocumentGenerator()
    generator.generate_all_test_documents()