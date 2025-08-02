#!/usr/bin/env python3
"""
Minimal OCR System with available dependencies
"""

import os
import sys
import time
import psutil
import numpy as np
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
import asyncio

# Resource Configuration
MAX_CPU_PERCENT = 70  # Maximum CPU usage percentage
MAX_RAM_PERCENT = 70  # Maximum RAM usage percentage  
MAX_RAM_GB = 2  # Maximum RAM in GB (can be exceeded if needed)

# Document processing imports
from PIL import Image
import cv2
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from fast_langdetect import detect_language
import pandas as pd

# Configure logging
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class OCRResult:
    """OCR Processing Result"""
    text: str
    language: str
    confidence: float
    tables: List[pd.DataFrame]
    markdown: str
    processing_time: float
    metadata: Dict[str, Any]


class ResourceManager:
    """Manages system resources to stay within limits"""
    
    def __init__(self):
        self.process = psutil.Process()
        self.start_time = time.time()
        
    def get_current_usage(self) -> Dict[str, float]:
        """Get current resource usage"""
        return {
            'cpu_percent': self.process.cpu_percent(interval=0.1),
            'ram_mb': self.process.memory_info().rss / (1024 * 1024),
            'ram_gb': self.process.memory_info().rss / (1024 * 1024 * 1024),
            'ram_percent': self.process.memory_percent()
        }
    
    def check_limits(self) -> bool:
        """Check if we're within resource limits"""
        usage = self.get_current_usage()
        
        if usage['cpu_percent'] > MAX_CPU_PERCENT:
            logger.warning(f"CPU usage {usage['cpu_percent']:.1f}% exceeds limit {MAX_CPU_PERCENT}%")
            
        if usage['ram_percent'] > MAX_RAM_PERCENT:
            logger.warning(f"RAM usage {usage['ram_percent']:.1f}% exceeds limit {MAX_RAM_PERCENT}%")
            
        return True
    
    def wait_for_resources(self, timeout: int = 10):
        """Wait for resources to become available"""
        start = time.time()
        while time.time() - start < timeout:
            if self.check_limits():
                return True
            time.sleep(0.5)
        return False


class MinimalOCREngine:
    """Minimal OCR Engine using available tools"""
    
    def __init__(self):
        self.resource_manager = ResourceManager()
        self._check_tesseract()
        
    def _check_tesseract(self):
        """Check if tesseract is available"""
        try:
            pytesseract.get_tesseract_version()
            logger.info("Tesseract is available")
        except:
            logger.warning("Tesseract not found. OCR functionality will be limited.")
            
    def detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            result = detect_language(text)
            return result['lang']
        except:
            return 'en'
            
    def extract_text_tesseract(self, image: np.ndarray) -> Tuple[str, float]:
        """Extract text using Tesseract"""
        try:
            # Convert to PIL Image
            pil_image = Image.fromarray(image)
            
            # Get text with confidence
            data = pytesseract.image_to_data(pil_image, output_type=pytesseract.Output.DICT)
            text = pytesseract.image_to_string(pil_image, lang='deu+eng')
            
            # Calculate confidence
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            confidence = np.mean(confidences) if confidences else 0
            
            return text, confidence / 100
        except Exception as e:
            logger.error(f"Tesseract OCR failed: {e}")
            return "", 0.0
        
    def process_image(self, image: np.ndarray) -> OCRResult:
        """Process a single image"""
        start_time = time.time()
        
        # Wait for resources
        self.resource_manager.wait_for_resources()
        
        # Preprocess image
        image = self._preprocess_image(image)
        
        # Extract text
        text, confidence = self.extract_text_tesseract(image)
        
        # Detect language
        language = self.detect_language(text) if text else 'unknown'
        
        # Convert to markdown
        markdown = text
        
        processing_time = time.time() - start_time
        
        return OCRResult(
            text=text,
            language=language,
            confidence=confidence,
            tables=[],
            markdown=markdown,
            processing_time=processing_time,
            metadata={
                'engine': 'tesseract',
                'image_shape': image.shape
            }
        )
        
    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR"""
        # Convert to grayscale if needed
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
            
        # Denoise
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Threshold
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        return thresh


class MinimalDocumentProcessor:
    """Process different document types with minimal dependencies"""
    
    def __init__(self, ocr_engine: MinimalOCREngine):
        self.ocr_engine = ocr_engine
        
    async def process_pdf(self, file_path: Path) -> List[OCRResult]:
        """Process PDF document"""
        results = []
        
        # Try native text extraction first
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            text = page.get_text()
            
            if text.strip():
                # Native text found
                results.append(OCRResult(
                    text=text,
                    language=self.ocr_engine.detect_language(text),
                    confidence=1.0,
                    tables=[],
                    markdown=text,
                    processing_time=0.1,
                    metadata={'page': page_num + 1, 'type': 'native'}
                ))
            else:
                # Convert to image and OCR
                pix = page.get_pixmap()
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                    pix.height, pix.width, pix.n
                )
                
                # Convert RGBA to RGB if needed
                if img.shape[2] == 4:
                    img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
                
                result = self.ocr_engine.process_image(img)
                result.metadata['page'] = page_num + 1
                results.append(result)
                
        doc.close()
        return results
        
    async def process_image(self, file_path: Path) -> OCRResult:
        """Process image file"""
        image = cv2.imread(str(file_path))
        return self.ocr_engine.process_image(image)
        
    async def process_docx(self, file_path: Path) -> OCRResult:
        """Process DOCX document"""
        doc = Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        return OCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=[],
            markdown=text,
            processing_time=0.1,
            metadata={'type': 'docx', 'paragraphs': len(doc.paragraphs)}
        )
        
    async def process_txt(self, file_path: Path) -> OCRResult:
        """Process text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        return OCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=[],
            markdown=text,
            processing_time=0.01,
            metadata={'type': 'txt', 'size': len(text)}
        )


class MinimalOCRSystem:
    """Minimal OCR System with basic functionality"""
    
    def __init__(self):
        self.ocr_engine = MinimalOCREngine()
        self.document_processor = MinimalDocumentProcessor(self.ocr_engine)
        
    async def process_file(self, file_path: str) -> OCRResult:
        """Process a single file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Determine file type and process
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            results = await self.document_processor.process_pdf(path)
            # Combine results from all pages
            combined_text = '\n\n'.join([r.text for r in results])
            combined_tables = []
            for r in results:
                combined_tables.extend(r.tables)
                
            return OCRResult(
                text=combined_text,
                language=results[0].language if results else 'en',
                confidence=np.mean([r.confidence for r in results]) if results else 0,
                tables=combined_tables,
                markdown='\n\n'.join([r.markdown for r in results]),
                processing_time=sum([r.processing_time for r in results]),
                metadata={'pages': len(results), 'type': 'pdf'}
            )
            
        elif suffix in ['.png', '.jpg', '.jpeg']:
            return await self.document_processor.process_image(path)
            
        elif suffix == '.docx':
            return await self.document_processor.process_docx(path)
            
        elif suffix == '.txt':
            return await self.document_processor.process_txt(path)
            
        else:
            raise ValueError(f"Unsupported file type: {suffix}")


def get_resource_info():
    """Get current resource usage information"""
    process = psutil.Process()
    return {
        'cpu_percent': process.cpu_percent(interval=1),
        'ram_mb': process.memory_info().rss / (1024 * 1024),
        'ram_gb': process.memory_info().rss / (1024 * 1024 * 1024),
        'ram_percent': process.memory_percent(),
        'threads': process.num_threads()
    }


def main():
    """Example usage"""
    print(f"Minimal OCR System Initialized")
    print(f"Resource Limits: CPU {MAX_CPU_PERCENT}%, RAM {MAX_RAM_PERCENT}% (Max {MAX_RAM_GB}GB)")
    
    # Show current resources
    resources = get_resource_info()
    print(f"\nCurrent Resources:")
    print(f"  CPU: {resources['cpu_percent']:.1f}%")
    print(f"  RAM: {resources['ram_gb']:.2f}GB ({resources['ram_percent']:.1f}%)")
    

if __name__ == "__main__":
    main()