#!/usr/bin/env python3
"""
Test Document Generator and Fetcher
Creates/downloads test documents in all supported formats with German text
"""

import os
import requests
import random
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import numpy as np
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import cv2
from bs4 import BeautifulSoup
import logging

logger = logging.getLogger(__name__)

# Test texts with German umlauts
GERMAN_TEXTS = [
    "Die schöne Müllerin ging über die Brücke.",
    "Äpfel, Öl und Übungen sind sehr nützlich.",
    "Der große Bär läuft durch den Wald.",
    "Frühstück mit Brötchen und Käse.",
    "Die Universität in München ist berühmt.",
    "Schöne Grüße aus Düsseldorf!",
    "Das Mädchen trägt ein blaues Kleid.",
    "Können Sie mir bitte helfen?",
    "Die Straße führt zum großen See.",
    "Süße Träume und gute Nacht!"
]

ENGLISH_TEXTS = [
    "The quick brown fox jumps over the lazy dog.",
    "Python is a powerful programming language.",
    "Machine learning revolutionizes data analysis.",
    "Welcome to the world of artificial intelligence.",
    "Document processing made easy with OCR.",
]

# Table data with umlauts
TABLE_DATA = [
    ["Stadt", "Einwohner", "Fläche (km²)"],
    ["München", "1,488,202", "310.7"],
    ["Köln", "1,083,498", "405.0"],
    ["Düsseldorf", "621,877", "217.4"],
    ["Nürnberg", "518,370", "186.5"],
    ["Würzburg", "127,880", "87.6"]
]


class TestDocumentGenerator:
    """Generate test documents in various formats"""
    
    def __init__(self, output_dir: str = "test_documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def generate_all(self):
        """Generate all test document types"""
        logger.info("Generating test documents...")
        
        # Generate text documents
        self.generate_txt_files()
        
        # Generate images
        self.generate_image_files()
        
        # Generate PDFs
        self.generate_pdf_files()
        
        # Generate DOCX
        self.generate_docx_files()
        
        # Download sample documents
        self.download_sample_documents()
        
        logger.info(f"Test documents generated in {self.output_dir}")
        
    def generate_txt_files(self):
        """Generate text files with German and English text"""
        # German text file
        with open(self.output_dir / "german_text.txt", "w", encoding="utf-8") as f:
            f.write("\n\n".join(GERMAN_TEXTS))
            
        # English text file
        with open(self.output_dir / "english_text.txt", "w", encoding="utf-8") as f:
            f.write("\n\n".join(ENGLISH_TEXTS))
            
        # Mixed language file
        with open(self.output_dir / "mixed_language.txt", "w", encoding="utf-8") as f:
            mixed = GERMAN_TEXTS[:5] + ENGLISH_TEXTS[:5]
            random.shuffle(mixed)
            f.write("\n\n".join(mixed))
            
    def generate_image_files(self):
        """Generate images with text"""
        # Create images with text
        for i, text in enumerate(GERMAN_TEXTS[:5]):
            self._create_text_image(
                text,
                self.output_dir / f"german_image_{i+1}.png"
            )
            
        for i, text in enumerate(ENGLISH_TEXTS[:3]):
            self._create_text_image(
                text,
                self.output_dir / f"english_image_{i+1}.jpg"
            )
            
        # Create a table image
        self._create_table_image(
            TABLE_DATA,
            self.output_dir / "table_image.png"
        )
        
        # Create a scanned-style image
        self._create_scanned_image(
            "\n".join(GERMAN_TEXTS),
            self.output_dir / "scanned_document.png"
        )
        
    def _create_text_image(self, text: str, output_path: Path):
        """Create an image with text"""
        # Create white background
        img = Image.new('RGB', (800, 200), color='white')
        draw = ImageDraw.Draw(img)
        
        # Try to use a font that supports umlauts
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", 36)
        except:
            font = ImageFont.load_default()
            
        # Draw text
        draw.text((50, 50), text, fill='black', font=font)
        
        # Save
        img.save(output_path)
        
    def _create_table_image(self, table_data: list, output_path: Path):
        """Create an image with a table"""
        # Calculate dimensions
        rows = len(table_data)
        cols = len(table_data[0])
        cell_width = 200
        cell_height = 50
        
        # Create image
        img = Image.new('RGB', 
                       (cols * cell_width + 100, rows * cell_height + 100), 
                       color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", 20)
        except:
            font = ImageFont.load_default()
            
        # Draw table
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                x = 50 + j * cell_width
                y = 50 + i * cell_height
                
                # Draw cell border
                draw.rectangle(
                    [x, y, x + cell_width, y + cell_height],
                    outline='black',
                    width=2
                )
                
                # Draw text
                draw.text((x + 10, y + 15), cell, fill='black', font=font)
                
        img.save(output_path)
        
    def _create_scanned_image(self, text: str, output_path: Path):
        """Create a scanned-style image with noise and rotation"""
        # Create base image
        img = Image.new('RGB', (1200, 1600), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", 24)
        except:
            font = ImageFont.load_default()
            
        # Draw text with line breaks
        y_position = 100
        for line in text.split('\n'):
            draw.text((100, y_position), line, fill='black', font=font)
            y_position += 40
            
        # Convert to numpy array
        img_array = np.array(img)
        
        # Add noise
        noise = np.random.normal(0, 10, img_array.shape).astype(np.uint8)
        img_array = cv2.add(img_array, noise)
        
        # Slight rotation
        rows, cols = img_array.shape[:2]
        angle = random.uniform(-2, 2)
        M = cv2.getRotationMatrix2D((cols/2, rows/2), angle, 1)
        img_array = cv2.warpAffine(img_array, M, (cols, rows), 
                                  borderValue=(255, 255, 255))
        
        # Add slight blur
        img_array = cv2.GaussianBlur(img_array, (3, 3), 0)
        
        # Save
        Image.fromarray(img_array).save(output_path)
        
    def generate_pdf_files(self):
        """Generate PDF files"""
        # Native PDF with text
        self._create_native_pdf(
            GERMAN_TEXTS,
            self.output_dir / "native_german.pdf"
        )
        
        # Native PDF with table
        self._create_pdf_with_table(
            TABLE_DATA,
            self.output_dir / "pdf_with_table.pdf"
        )
        
        # Scanned PDF (images embedded)
        self._create_scanned_pdf(
            self.output_dir / "scanned_german.pdf"
        )
        
    def _create_native_pdf(self, texts: list, output_path: Path):
        """Create a native PDF with searchable text"""
        pdf = FPDF()
        pdf.add_page()
        
        # Add Unicode font support
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', size=12)
        
        # Add title
        pdf.cell(0, 10, 'German Text Document', ln=True, align='C')
        pdf.ln(10)
        
        # Add texts
        for text in texts:
            pdf.multi_cell(0, 10, text)
            pdf.ln(5)
            
        pdf.output(str(output_path))
        
    def _create_pdf_with_table(self, table_data: list, output_path: Path):
        """Create PDF with table"""
        pdf = FPDF()
        pdf.add_page()
        
        # Add Unicode font
        pdf.add_font('DejaVu', '', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', size=12)
        
        # Title
        pdf.cell(0, 10, 'German Cities Data', ln=True, align='C')
        pdf.ln(10)
        
        # Table
        col_width = 60
        row_height = 10
        
        # Header
        pdf.set_font('DejaVu', size=12)
        for item in table_data[0]:
            pdf.cell(col_width, row_height, item, border=1)
        pdf.ln(row_height)
        
        # Data
        pdf.set_font('DejaVu', size=11)
        for row in table_data[1:]:
            for item in row:
                pdf.cell(col_width, row_height, item, border=1)
            pdf.ln(row_height)
            
        pdf.output(str(output_path))
        
    def _create_scanned_pdf(self, output_path: Path):
        """Create a scanned PDF (images only)"""
        # First create some images
        temp_images = []
        
        for i in range(3):
            img_path = self.output_dir / f"temp_scan_{i}.png"
            text = "\n".join(GERMAN_TEXTS[i*3:(i+1)*3])
            self._create_scanned_image(text, img_path)
            temp_images.append(img_path)
            
        # Convert images to PDF
        images = [Image.open(img) for img in temp_images]
        
        # Save as PDF
        images[0].save(
            output_path,
            save_all=True,
            append_images=images[1:]
        )
        
        # Clean up temp images
        for img_path in temp_images:
            img_path.unlink()
            
    def generate_docx_files(self):
        """Generate DOCX files"""
        # Simple DOCX
        self._create_simple_docx(
            GERMAN_TEXTS,
            self.output_dir / "german_document.docx"
        )
        
        # DOCX with table
        self._create_docx_with_table(
            TABLE_DATA,
            self.output_dir / "docx_with_table.docx"
        )
        
    def _create_simple_docx(self, texts: list, output_path: Path):
        """Create simple DOCX document"""
        doc = Document()
        
        # Add title
        doc.add_heading('German Text Document', 0)
        
        # Add paragraphs
        for text in texts:
            doc.add_paragraph(text)
            
        doc.save(str(output_path))
        
    def _create_docx_with_table(self, table_data: list, output_path: Path):
        """Create DOCX with table"""
        doc = Document()
        
        # Add title
        doc.add_heading('German Cities Data', 0)
        doc.add_paragraph('')
        
        # Add table
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Fill table
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
                
        doc.save(str(output_path))
        
    def download_sample_documents(self):
        """Download sample documents from the internet"""
        samples = [
            {
                'url': 'https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf',
                'name': 'sample_english.pdf'
            },
            {
                'url': 'https://upload.wikimedia.org/wikipedia/commons/thumb/b/b6/Gutenberg_Bible%2C_Lenox_Copy%2C_New_York_Public_Library%2C_2009._Pic_01.jpg/1200px-Gutenberg_Bible%2C_Lenox_Copy%2C_New_York_Public_Library%2C_2009._Pic_01.jpg',
                'name': 'historical_scan.jpg'
            }
        ]
        
        for sample in samples:
            try:
                response = requests.get(sample['url'], timeout=10)
                if response.status_code == 200:
                    with open(self.output_dir / sample['name'], 'wb') as f:
                        f.write(response.content)
                    logger.info(f"Downloaded {sample['name']}")
            except Exception as e:
                logger.error(f"Failed to download {sample['name']}: {e}")


def main():
    """Generate all test documents"""
    generator = TestDocumentGenerator()
    generator.generate_all()
    
    # List generated files
    print("\nGenerated test documents:")
    for file in sorted(generator.output_dir.iterdir()):
        if file.is_file():
            size_kb = file.stat().st_size / 1024
            print(f"  - {file.name} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()