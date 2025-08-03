#!/usr/bin/env python3
"""
CPU-Only OCR System - Optimized for CPU usage without GPU dependencies
"""

import os
import sys
import time
import gc
import psutil
import numpy as np
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
from dataclasses import dataclass, field
import asyncio
import re
import multiprocessing as mp
from concurrent.futures import ThreadPoolExecutor

# CPU-Only Resource Configuration
# Maximized for full resource utilization with safety margins
MAX_CPU_PERCENT = 95  # Maximum CPU usage percentage (95% to leave headroom for system)
MAX_RAM_PERCENT = 90  # Maximum RAM usage percentage (90% to prevent OOM)
MAX_RAM_GB = 8  # Maximum RAM in GB (increased for better performance)
MAX_WORKERS = mp.cpu_count() if mp.cpu_count() <= 8 else 8  # Use all cores, cap at 8

# Document processing imports
from PIL import Image
import cv2
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from docx import Document
import pytesseract
from fast_langdetect import detect_language
import pandas as pd
import pdfplumber

# Configure logging
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class CPUOptimizedConfig:
    """CPU-optimized configuration"""
    use_gpu: bool = False  # Always False for CPU-only
    max_workers: int = MAX_WORKERS
    tesseract_threads: int = 4  # Tesseract can use multiple threads
    batch_size: int = 5  # Smaller batches for CPU
    enable_preprocessing: bool = True  # Image preprocessing helps accuracy
    dpi_for_pdf: int = 200  # Lower DPI for faster processing
    jpeg_quality: int = 85  # Balance between quality and speed


@dataclass
class TableData:
    """Represents extracted table data"""
    headers: List[str]
    rows: List[List[str]]
    confidence: float = 0.0
    position: Dict[str, int] = field(default_factory=dict)
    
    def to_dataframe(self) -> pd.DataFrame:
        """Convert to pandas DataFrame"""
        if self.headers:
            return pd.DataFrame(self.rows, columns=self.headers)
        else:
            return pd.DataFrame(self.rows)
            
    def to_markdown(self) -> str:
        """Convert to markdown table"""
        if not self.rows:
            return ""
            
        headers = self.headers if self.headers else [f"Col{i+1}" for i in range(len(self.rows[0]))]
        
        lines = []
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join([" --- " for _ in headers]) + "|")
        
        for row in self.rows:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
            
        return "\n".join(lines)


@dataclass
class CPUOCRResult:
    """OCR Result optimized for CPU processing"""
    text: str
    language: str
    confidence: float
    tables: List[TableData]
    markdown: str
    processing_time: float
    metadata: Dict[str, Any]
    cpu_usage: float = 0.0
    ram_usage_mb: float = 0.0


class CPUResourceManager:
    """CPU-specific resource management"""
    
    def __init__(self):
        self.process = psutil.Process()
        self.start_time = time.time()
        self._initial_cpu_affinity = None
        
    def limit_cpu_cores(self, max_cores: int = None):
        """Limit process to specific number of CPU cores"""
        try:
            if max_cores is None:
                max_cores = MAX_WORKERS
                
            available_cores = list(range(mp.cpu_count()))
            if max_cores < len(available_cores):
                # Use first N cores
                self.process.cpu_affinity(available_cores[:max_cores])
                logger.info(f"Limited to {max_cores} CPU cores")
        except Exception as e:
            logger.warning(f"Could not set CPU affinity: {e}")
            
    def get_current_usage(self) -> Dict[str, float]:
        """Get current resource usage"""
        cpu_percent = self.process.cpu_percent(interval=0.1)
        memory_info = self.process.memory_info()
        
        return {
            'cpu_percent': cpu_percent,
            'cpu_count': mp.cpu_count(),
            'ram_mb': memory_info.rss / (1024 * 1024),
            'ram_gb': memory_info.rss / (1024 * 1024 * 1024),
            'ram_percent': self.process.memory_percent(),
            'threads': self.process.num_threads()
        }
        
    def wait_for_resources(self, timeout: int = 10) -> bool:
        """Wait for CPU resources to become available"""
        start = time.time()
        
        while time.time() - start < timeout:
            usage = self.get_current_usage()
            
            # Check CPU usage with dynamic threshold
            if usage['cpu_percent'] < MAX_CPU_PERCENT:
                return True
                
            # If CPU is busy, wait with exponential backoff
            wait_time = min(0.1 * (2 ** ((time.time() - start) / 2)), 2.0)
            time.sleep(wait_time)
            
            # Force garbage collection if RAM is high
            if usage['ram_percent'] > 85:
                import gc
                gc.collect()
                logger.info(f"Triggered GC: RAM usage {usage['ram_percent']:.1f}%")
            
        logger.warning("Timeout waiting for CPU resources")
        return False
        
    def optimize_for_cpu(self):
        """Apply CPU-specific optimizations"""
        # Set process priority to normal (not high)
        try:
            if sys.platform == 'linux':
                os.nice(10)  # Lower priority
            logger.info("Set process to lower priority for better system responsiveness")
        except:
            pass
            
        # Limit CPU cores if needed
        self.limit_cpu_cores()


import gc

class CPUOptimizedOCREngine:
    """OCR Engine optimized for CPU-only processing"""
    
    def __init__(self, config: CPUOptimizedConfig = None):
        self.config = config or CPUOptimizedConfig()
        self.resource_manager = CPUResourceManager()
        self.resource_manager.optimize_for_cpu()
        self._check_tesseract()
        
        # Thread pool for parallel processing
        self.executor = ThreadPoolExecutor(max_workers=self.config.max_workers)
        
    def _check_tesseract(self):
        """Check if tesseract is available and get version"""
        try:
            version = pytesseract.get_tesseract_version()
            logger.info(f"Tesseract {version} is available")
            
            # Check for German language data
            langs = pytesseract.get_languages(config='')
            if 'deu' in langs:
                logger.info("German language data available")
            else:
                logger.warning("German language data not found")
                
        except Exception as e:
            logger.error(f"Tesseract not found: {e}")
            raise RuntimeError("Tesseract is required for CPU-only OCR")
            
    def detect_language(self, text: str) -> str:
        """Detect language of text"""
        try:
            if len(text.strip()) > 10:
                result = detect_language(text)
                return result['lang']
        except:
            pass
        return 'de'
        
    def _preprocess_image_cpu_optimized(self, image: np.ndarray) -> np.ndarray:
        """CPU-optimized image preprocessing - gentle for better Umlaut recognition"""
        # Convert to grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
            
        # Check image quality
        mean_brightness = np.mean(gray)
        std_brightness = np.std(gray)
        
        # If image has good contrast and brightness, return as is
        if 80 < mean_brightness < 220 and std_brightness > 30:
            return gray
            
        # Apply gentle contrast enhancement
        # Avoid binary thresholding which can destroy Umlauts
        if mean_brightness < 100:
            # Dark image - brighten it
            adjusted = cv2.convertScaleAbs(gray, alpha=1.3, beta=30)
        elif mean_brightness > 200:
            # Bright image - increase contrast
            adjusted = cv2.convertScaleAbs(gray, alpha=1.2, beta=-30)
        else:
            # Normal brightness - just enhance contrast slightly
            clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8, 8))
            adjusted = clahe.apply(gray)
        
        return adjusted
        
    def extract_text_cpu(self, image: np.ndarray) -> Tuple[str, float, Dict[str, Any]]:
        """CPU-optimized text extraction using Tesseract"""
        start_time = time.time()
        
        # Preprocess if enabled
        if self.config.enable_preprocessing:
            processed_image = self._preprocess_image_cpu_optimized(image)
        else:
            processed_image = image
            
        # Convert to PIL
        pil_image = Image.fromarray(processed_image)
        
        # Configure Tesseract for CPU optimization
        custom_config = r'--oem 3 --psm 6'  # Use default OCR engine mode
        
        # Get detailed data
        data = pytesseract.image_to_data(
            pil_image, 
            output_type=pytesseract.Output.DICT,
            lang='deu+eng',
            config=custom_config
        )
        
        # Extract text
        text = pytesseract.image_to_string(
            pil_image,
            lang='deu+eng',
            config=custom_config
        )
        
        # Calculate confidence
        confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
        confidence = np.mean(confidences) if confidences else 0
        
        processing_time = time.time() - start_time
        
        return text, confidence / 100, {
            'processing_time': processing_time,
            'words': len([w for w in data['text'] if w.strip()]),
            'method': 'tesseract_cpu'
        }
        
    def process_image_cpu_optimized(self, image: np.ndarray) -> CPUOCRResult:
        """Process image with CPU optimization and memory management"""
        start_time = time.time()
        initial_usage = self.resource_manager.get_current_usage()
        
        # Check memory before processing
        if initial_usage['ram_percent'] > 85:
            logger.warning(f"High memory usage: {initial_usage['ram_percent']:.1f}%")
            gc.collect()  # Force garbage collection
            # If still high, reduce image size
            if image.shape[0] > 2000 or image.shape[1] > 2000:
                scale = 0.7
                new_shape = (int(image.shape[1] * scale), int(image.shape[0] * scale))
                image = cv2.resize(image, new_shape, interpolation=cv2.INTER_AREA)
                logger.info(f"Reduced image size to {new_shape} due to memory constraints")
        
        # Wait for resources
        self.resource_manager.wait_for_resources()
        
        try:
            # Extract text
            text, confidence, metadata = self.extract_text_cpu(image)
            
            # Extract tables (simplified for CPU)
            tables = self._extract_tables_cpu(image, text)
            
            # Detect language
            language = self.detect_language(text)
            
            # Generate markdown
            markdown = self._generate_markdown(text, tables)
            
        except Exception as e:
            logger.error(f"Error during OCR processing: {e}")
            # Return partial result on error
            text = "Error during OCR processing"
            confidence = 0.0
            tables = []
            language = "unknown"
            markdown = text
            metadata = {'error': str(e)}
        finally:
            # Clean up to free memory
            del image
            gc.collect()
        
        # Get final resource usage
        final_usage = self.resource_manager.get_current_usage()
        
        processing_time = time.time() - start_time
        
        return CPUOCRResult(
            text=text,
            language=language,
            confidence=confidence,
            tables=tables,
            markdown=markdown,
            processing_time=processing_time,
            metadata=metadata,
            cpu_usage=final_usage['cpu_percent'],
            ram_usage_mb=final_usage['ram_mb']
        )
        
    def _extract_tables_cpu(self, image: np.ndarray, text: str) -> List[TableData]:
        """Simplified table extraction for CPU"""
        tables = []
        
        # Try to extract tables from text structure
        lines = text.split('\n')
        potential_tables = self._find_text_tables(lines)
        
        for table_lines in potential_tables:
            table_data = self._parse_text_table(table_lines)
            if table_data:
                tables.append(table_data)
                
        return tables
        
    def _find_text_tables(self, lines: List[str]) -> List[List[str]]:
        """Find potential tables in text lines"""
        potential_tables = []
        current_table = []
        
        for line in lines:
            # Simple heuristic: line with multiple columns
            if self._is_table_line(line):
                current_table.append(line)
            else:
                if len(current_table) >= 2:  # At least 2 rows
                    potential_tables.append(current_table)
                current_table = []
                
        if len(current_table) >= 2:
            potential_tables.append(current_table)
            
        return potential_tables
        
    def _is_table_line(self, line: str) -> bool:
        """Check if line looks like a table row"""
        # Count separators
        separators = len(re.findall(r'\t|\s{2,}|\|', line))
        return separators >= 1 and len(line.strip()) > 0
        
    def _parse_text_table(self, lines: List[str]) -> Optional[TableData]:
        """Parse text lines as table"""
        if not lines:
            return None
            
        rows = []
        for line in lines:
            # Try different separators
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            elif '\t' in line:
                cells = [cell.strip() for cell in line.split('\t')]
            else:
                cells = re.split(r'\s{2,}', line.strip())
                
            if len(cells) >= 2:
                rows.append(cells)
                
        if rows:
            return TableData(
                headers=rows[0] if len(rows) > 1 else [],
                rows=rows[1:] if len(rows) > 1 else rows,
                confidence=0.7
            )
            
        return None
        
    def _generate_markdown(self, text: str, tables: List[TableData]) -> str:
        """Generate markdown output"""
        parts = []
        
        if text:
            parts.append(text)
            
        if tables:
            parts.append("\n## Extracted Tables\n")
            for i, table in enumerate(tables):
                parts.append(f"\n### Table {i+1}")
                parts.append(table.to_markdown())
                
        return "\n".join(parts)


class CPUDocumentProcessor:
    """Document processor optimized for CPU"""
    
    def __init__(self, ocr_engine: CPUOptimizedOCREngine):
        self.ocr_engine = ocr_engine
        self.resource_manager = ocr_engine.resource_manager
        
    async def process_pdf_cpu(self, file_path: Path) -> List[CPUOCRResult]:
        """Process PDF with CPU optimization"""
        results = []
        
        # Try pdfplumber first for text extraction
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    start_time = time.time()
                    
                    # Extract text
                    text = page.extract_text() or ""
                    
                    # Extract tables
                    tables = []
                    for table in page.extract_tables():
                        if table and len(table) >= 2:
                            table_data = TableData(
                                headers=table[0] if table else [],
                                rows=table[1:] if len(table) > 1 else [],
                                confidence=0.95
                            )
                            tables.append(table_data)
                            
                    if text or tables:
                        result = CPUOCRResult(
                            text=text,
                            language=self.ocr_engine.detect_language(text),
                            confidence=1.0,
                            tables=tables,
                            markdown=self.ocr_engine._generate_markdown(text, tables),
                            processing_time=time.time() - start_time,
                            metadata={'page': page_num + 1, 'method': 'pdfplumber'},
                            cpu_usage=0,
                            ram_usage_mb=0
                        )
                        results.append(result)
                        
        except Exception as e:
            logger.debug(f"pdfplumber failed, using PyMuPDF: {e}")
            
            # Fallback to PyMuPDF
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                
                if text.strip():
                    # Native text found
                    result = CPUOCRResult(
                        text=text,
                        language=self.ocr_engine.detect_language(text),
                        confidence=1.0,
                        tables=[],
                        markdown=text,
                        processing_time=0.1,
                        metadata={'page': page_num + 1, 'method': 'native'},
                        cpu_usage=0,
                        ram_usage_mb=0
                    )
                    results.append(result)
                else:
                    # Convert to image for OCR
                    # Use lower DPI for CPU efficiency
                    pix = page.get_pixmap(dpi=self.ocr_engine.config.dpi_for_pdf)
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                        pix.height, pix.width, pix.n
                    )
                    
                    if img.shape[2] == 4:
                        img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
                        
                    result = self.ocr_engine.process_image_cpu_optimized(img)
                    result.metadata['page'] = page_num + 1
                    results.append(result)
                    
            doc.close()
            
        return results
        
    async def process_image_cpu(self, file_path: Path) -> CPUOCRResult:
        """Process image file with CPU optimization"""
        # Read image
        image = cv2.imread(str(file_path))
        
        # Dynamic resize based on available memory
        usage = self.resource_manager.get_current_usage()
        if usage['ram_percent'] > 70:
            max_dimension = 1500  # More aggressive resize if memory is high
        else:
            max_dimension = 3000  # Allow larger images when memory is available
            
        height, width = image.shape[:2]
        
        if width > max_dimension or height > max_dimension:
            scale = max_dimension / max(width, height)
            new_width = int(width * scale)
            new_height = int(height * scale)
            image = cv2.resize(image, (new_width, new_height), interpolation=cv2.INTER_AREA)
            logger.info(f"Resized image from {width}x{height} to {new_width}x{new_height} (RAM: {usage['ram_percent']:.1f}%)")
            
        return self.ocr_engine.process_image_cpu_optimized(image)
        
    async def process_docx_cpu(self, file_path: Path) -> CPUOCRResult:
        """Process DOCX file"""
        start_time = time.time()
        doc = Document(file_path)
        
        # Extract text
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
                
        text = '\n'.join(paragraphs)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
                    
            if rows:
                table_data = TableData(
                    headers=rows[0] if len(rows) > 1 else [],
                    rows=rows[1:] if len(rows) > 1 else rows,
                    confidence=1.0
                )
                tables.append(table_data)
                
        return CPUOCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=tables,
            markdown=self.ocr_engine._generate_markdown(text, tables),
            processing_time=time.time() - start_time,
            metadata={'type': 'docx', 'paragraphs': len(paragraphs)},
            cpu_usage=0,
            ram_usage_mb=0
        )
        
    async def process_rtf_cpu(self, file_path: Path) -> CPUOCRResult:
        """Process RTF file"""
        start_time = time.time()
        
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            text = rtf_to_text(rtf_content)
        except Exception as e:
            logger.error(f"Error processing RTF file: {e}")
            # Fallback: try to read as plain text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        # Extract tables (RTF tables are harder to parse, so we use simple heuristics)
        tables = self.ocr_engine._extract_tables_cpu(None, text)
        
        return CPUOCRResult(
            text=text,
            language=self.ocr_engine.detect_language(text),
            confidence=1.0,
            tables=tables,
            markdown=self.ocr_engine._generate_markdown(text, tables),
            processing_time=time.time() - start_time,
            metadata={'type': 'rtf', 'size': len(text)},
            cpu_usage=0,
            ram_usage_mb=0
        )
        
    async def process_odt_cpu(self, file_path: Path) -> CPUOCRResult:
        """Process ODT (OpenDocument Text) file"""
        start_time = time.time()
        
        try:
            from odf import text, teletype
            from odf.opendocument import load
            
            doc = load(str(file_path))
            paragraphs = []
            
            # Extract paragraphs
            for p in doc.getElementsByType(text.P):
                text_content = teletype.extractText(p)
                if text_content.strip():
                    paragraphs.append(text_content)
            
            extracted_text = '\n'.join(paragraphs)
            
            # Extract tables
            tables = []
            from odf.table import Table, TableRow, TableCell
            for table in doc.getElementsByType(Table):
                rows = []
                for tr in table.getElementsByType(TableRow):
                    cells = []
                    for td in tr.getElementsByType(TableCell):
                        cell_text = teletype.extractText(td).strip()
                        cells.append(cell_text)
                    if any(cells):
                        rows.append(cells)
                
                if rows:
                    table_data = TableData(
                        headers=rows[0] if len(rows) > 1 else [],
                        rows=rows[1:] if len(rows) > 1 else rows,
                        confidence=1.0
                    )
                    tables.append(table_data)
                    
        except Exception as e:
            logger.error(f"Error processing ODT file: {e}")
            # Fallback: try to extract text using a simpler method
            import zipfile
            import xml.etree.ElementTree as ET
            
            extracted_text = ""
            try:
                with zipfile.ZipFile(file_path, 'r') as z:
                    with z.open('content.xml') as f:
                        tree = ET.parse(f)
                        root = tree.getroot()
                        
                        # Extract all text nodes
                        texts = []
                        for elem in root.iter():
                            if elem.text:
                                texts.append(elem.text)
                        extracted_text = ' '.join(texts)
            except:
                extracted_text = "Error extracting ODT content"
            
            tables = []
        
        return CPUOCRResult(
            text=extracted_text,
            language=self.ocr_engine.detect_language(extracted_text),
            confidence=1.0,
            tables=tables,
            markdown=self.ocr_engine._generate_markdown(extracted_text, tables),
            processing_time=time.time() - start_time,
            metadata={'type': 'odt', 'paragraphs': len(paragraphs) if 'paragraphs' in locals() else 0},
            cpu_usage=0,
            ram_usage_mb=0
        )


class CPUOnlyOCRSystem:
    """Complete CPU-only OCR System"""
    
    def __init__(self, config: CPUOptimizedConfig = None):
        self.config = config or CPUOptimizedConfig()
        self.resource_manager = CPUResourceManager()
        self.ocr_engine = CPUOptimizedOCREngine(self.config)
        self.document_processor = CPUDocumentProcessor(self.ocr_engine)
        
        logger.info(f"CPU-Only OCR System initialized")
        logger.info(f"Max workers: {self.config.max_workers}")
        logger.info(f"CPU cores available: {mp.cpu_count()}")
        
    async def process_file(self, file_path: str) -> CPUOCRResult:
        """Process file with CPU optimization and resource management"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Check available resources before processing
        if not self.resource_manager.wait_for_resources(timeout=30):
            logger.warning("Resource limits reached, proceeding anyway for maximum utilization")
            
        suffix = path.suffix.lower()
        
        logger.info(f"Processing {suffix} file: {path.name}")
        
        # Monitor resources during processing
        initial_usage = self.resource_manager.get_current_usage()
        logger.info(f"Initial resources - CPU: {initial_usage['cpu_percent']:.1f}%, RAM: {initial_usage['ram_percent']:.1f}%")
        
        if suffix == '.pdf':
            results = await self.document_processor.process_pdf_cpu(path)
            
            if not results:
                return self._empty_result()
                
            # Combine results
            combined_text = '\n\n'.join([r.text for r in results])
            all_tables = []
            for r in results:
                all_tables.extend(r.tables)
                
            return CPUOCRResult(
                text=combined_text,
                language=results[0].language if results else 'de',
                confidence=np.mean([r.confidence for r in results]) if results else 0,
                tables=all_tables,
                markdown=self.ocr_engine._generate_markdown(combined_text, all_tables),
                processing_time=sum([r.processing_time for r in results]),
                metadata={
                    'pages': len(results),
                    'type': 'pdf',
                    'total_tables': len(all_tables)
                },
                cpu_usage=max([r.cpu_usage for r in results]) if results else 0,
                ram_usage_mb=max([r.ram_usage_mb for r in results]) if results else 0
            )
            
        elif suffix in ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif', '.webp']:
            return await self.document_processor.process_image_cpu(path)
            
        elif suffix == '.docx':
            return await self.document_processor.process_docx_cpu(path)
            
        elif suffix in ['.txt', '.csv', '.xml', '.html', '.htm']:
            # Process text-based files
            encoding = 'utf-8'
            try:
                with open(path, 'r', encoding=encoding) as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Try different encodings
                for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(path, 'r', encoding=enc) as f:
                            text = f.read()
                        encoding = enc
                        break
                    except:
                        continue
                else:
                    raise ValueError(f"Unable to decode file {path.name}")
            
            # For CSV files, format as table
            if suffix == '.csv':
                try:
                    import pandas as pd
                    df = pd.read_csv(path, encoding=encoding)
                    text = df.to_string()
                except:
                    pass  # Keep raw text if CSV parsing fails
                
            # Try to extract tables from text
            tables = self.ocr_engine._extract_tables_cpu(None, text)
            
            return CPUOCRResult(
                text=text,
                language=self.ocr_engine.detect_language(text),
                confidence=1.0,
                tables=tables,
                markdown=self.ocr_engine._generate_markdown(text, tables),
                processing_time=0.01,
                metadata={'type': suffix[1:], 'size': len(text), 'encoding': encoding},
                cpu_usage=0,
                ram_usage_mb=0
            )
            
        elif suffix == '.rtf':
            return await self.document_processor.process_rtf_cpu(path)
            
        elif suffix == '.odt':
            return await self.document_processor.process_odt_cpu(path)
            
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
            
    def _empty_result(self) -> CPUOCRResult:
        """Return empty result"""
        return CPUOCRResult(
            text="",
            language="unknown",
            confidence=0.0,
            tables=[],
            markdown="",
            processing_time=0.0,
            metadata={},
            cpu_usage=0,
            ram_usage_mb=0
        )
        
    async def process_batch(self, file_paths: List[str]) -> List[CPUOCRResult]:
        """Process multiple files with optimized resource utilization"""
        results = []
        
        # Process files with parallel workers for maximum CPU utilization
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            # Submit all files for processing
            future_to_file = {executor.submit(asyncio.run, self.process_file(fp)): fp 
                             for fp in file_paths}
            
            for future in concurrent.futures.as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    result = future.result()
                    results.append(result)
                    logger.info(f"Completed processing: {file_path}")
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    # Add empty result to maintain order
                    results.append(self._empty_result())
                
        return results


def get_resource_info():
    """Get current system resource information"""
    process = psutil.Process()
    return {
        'cpu_percent': process.cpu_percent(interval=1),
        'cpu_count': mp.cpu_count(),
        'ram_mb': process.memory_info().rss / (1024 * 1024),
        'ram_gb': process.memory_info().rss / (1024 * 1024 * 1024),
        'ram_percent': process.memory_percent(),
        'threads': process.num_threads()
    }


def main():
    """Test CPU-only OCR system"""
    print("CPU-Only OCR System")
    print(f"Resource Limits: CPU {MAX_CPU_PERCENT}%, RAM {MAX_RAM_PERCENT}% (Max {MAX_RAM_GB}GB)")
    print(f"CPU Cores: {mp.cpu_count()} available, using max {MAX_WORKERS}")
    print("\nFeatures:")
    print("  - Tesseract-based OCR (CPU-optimized)")
    print("  - Automatic image preprocessing")
    print("  - Table extraction from text structure")
    print("  - Resource management for CPU")
    print("  - Batch processing support")
    
    # Show current resources
    resources = get_resource_info()
    print(f"\nCurrent Resources:")
    print(f"  CPU: {resources['cpu_percent']:.1f}%")
    print(f"  RAM: {resources['ram_gb']:.2f}GB ({resources['ram_percent']:.1f}%)")
    print(f"  Threads: {resources['threads']}")


if __name__ == "__main__":
    main()